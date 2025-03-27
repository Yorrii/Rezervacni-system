import hashlib
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from flask_login import UserMixin
from datetime import datetime


def porovnat_hesla(heslo, db_heslo):
    """
    Metoda slouží k porovnání hesel

    Parametry:
        heslo (str): heslo, které uživatel zadá při přihlašování na webové stránce
        db_heslo (str): heslo, které se vezme z databáze podle zadaného emailu.

    Vrací:
        True (boolean): pokud jsou oba stringy stejné
        False (boolean): pokud jsou stringy odlišné
    """
    return True if hashlib.sha256(heslo.encode('utf-8')).hexdigest() == db_heslo else False

class User(UserMixin):
    def __init__(self, id, isAdmin=False):
        self.id = id
        if int(self.id) > 1000000: # pokud je superadmin
            self.isSuperAdmin = True
            self.isCommissar = True
            self.isAdmin = True
        elif int(self.id) > 100000: # pokud je komisař
            self.isSuperAdmin = False 
            self.isCommissar = True
            self.isAdmin = isAdmin # pokud má admin práva
        else: # pro autoškolu
            self.isSuperAdmin = False
            self.isAdmin = False
            self.isCommissar = False

def create_student_document(datum, seznam_studentu):
    """
    Vytvoří profesionálně vypadající Word dokument se seznamem žáků pro daný termín.

    Parametry:
        datum (str): Datum termínu ve formátu 'DD.MM.YYYY'.
        seznam_studentu (list): Seznam objektů žáků, kteří jsou přihlášeni k termínu.

    Návratová hodnota:
        Document: Objekt Word dokumentu s tabulkou obsahující informace o žácích.
    """
    doc = Document()

    # Přidání nadpisu
    title = doc.add_paragraph()
    title_run = title.add_run(f"Seznam žáků pro termín {datum}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Zarovnání na střed

    doc.add_paragraph()  # Přidání mezery pod nadpisem

    # Vytvoření tabulky se stylem
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Zarovnání tabulky na střed

    # Definování nadpisů sloupců
    headers = [
        "Evidenční číslo", "Jméno", "Příjmení",
        "Začátek výuky", "Datum narození", "Typ zkoušky", "Druh zkoušky"
    ]
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Zarovnání na střed
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(11)  # Zvětšení fontu

    # Vyplnění tabulky daty studentů
    for student in seznam_studentu:
        row_cells = table.add_row().cells
        row_cells[0].text = student.zak.ev_cislo
        row_cells[1].text = student.zak.jmeno
        row_cells[2].text = student.zak.prijmeni
        row_cells[3].text = _format_date(student.zak.zacatek)
        row_cells[4].text = _format_date(student.zak.narozeni)
        row_cells[5].text = student.typ_zkousky
        row_cells[6].text = student.druh_zkousky

        # Zarovnání obsahu buněk a nastavení velikosti písma
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    return doc

def _format_date(date):
    """Pomocná funkce pro formátování data."""
    return date.strftime("%d.%m.%Y") if date else "Neznámé"


def create_document(autoskola, datum, seznam_studentu, komisar): 
    doc = Document()
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    #doc.sections[0].right_margin = 457200
    #doc.sections[0].left_margin = 457200
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(1.4)

    #header
    header = doc.sections[0].header
    
    table = header.add_table(rows=1, cols=2, width=Pt(595))

    table.columns[0].width = Pt(295)
    table.columns[1].width = Pt(290)
    # záhlavý logo
    cell_logo = table.cell(0, 0)
    cell_logo_paragraph = cell_logo.paragraphs[0]
    run = cell_logo_paragraph.add_run()
    run.add_picture('Code/static/images/erb.png', width=Pt(60))

    # záhlavý pravá část
    cell_right = table.cell(0, 1)
    paragraf = cell_right.paragraphs[0]
    run = paragraf.add_run("MmM Z_OSČ_303")
    run.font.size = Pt(8)
    run.font.name = 'Work Sans'
    cell_right_paragraph = cell_right.paragraphs[0]
    right_text = "\n\nMAGISTRÁT MĚSTA MOSTU\nODBOR SPRÁVNÍCH ČINNOSTÍ"
    run = cell_right_paragraph.add_run(right_text)
    run.font.name = 'Work Sans'
    run.font.size = Pt(10)
    run.bold = True
    run = cell_right_paragraph.add_run("\noddělení registru řidičů a vozidel")
    run.font.name = 'Work Sans'
    run.font.size = Pt(10)
    cell_right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell_right_paragraph.paragraph_format.space_after = Pt(0)
    #BODY
    # info o komisaři
    part2= doc.add_table(rows=1, cols=2)
    part2.columns[0].width = Pt(250)
    part2.columns[1].width = Pt(250)

    part2_left = part2.cell(0, 0)
    def add_paragraph_with_spacing(text, font_size=9):
        paragraph = part2_left.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        paragraph.paragraph_format.space_after = Pt(0)  # Nastavení mezer po odstavci
        paragraph.paragraph_format.space_before = Pt(0) # Nastavení mezer před odstavcem
        return paragraph
    
    add_paragraph_with_spacing(f"Váš dopis zn.:")
    add_paragraph_with_spacing(f"Naše zn.:")
    add_paragraph_with_spacing(f"Listů/příloh: 0/0")
    add_paragraph_with_spacing(f"Vyřizuje: {komisar.jmeno} {komisar.prijmeni}")
    add_paragraph_with_spacing(f"Telefon:")
    add_paragraph_with_spacing(f"Email: {komisar.email}")
    add_paragraph_with_spacing(f"Most, {datetime.today().strftime("%d.%m.%Y")}")
    # info o autoškole
    part2_right = part2.cell(0, 1)

    # Vytvoř nový odstavec v pravé buňce a nastav zarovnání doprava
    adresaFormated = autoskola.adresa_u.split(', ')
    part2_right_paragraph = part2_right.add_paragraph()
    part2_right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    part2_text = f"{autoskola.nazev}\n{adresaFormated[0]}\n{adresaFormated[1]}\n{adresaFormated[2]}"   
    run = part2_right_paragraph.add_run(part2_text)
    run.bold = False
    run.font.name = 'Arial'
    run.font.size = Pt(12)

# Zarovnej každý řádek v pravé buňce na levý okraj
    for paragraph in part2_right.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph()
    paragraf = doc.add_paragraph()
    run = paragraf.add_run('Zkouška z odborné způsobilosti k řízení motorového vozidla')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.bold = True
    
    paragraf = doc.add_paragraph()
    paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = paragraf.add_run("""\nNa základě Vámi podané přihlášky k provedení zkoušky níže uvedeného uchazeče o řidičské oprávnění Vám sdělujeme, že tato zkouška bude vykonána v souladu s ustanovením §32 a §39 zákona č. 247/2000 Sb. \no získávání a zdokonalování odborné způsobilosti k řízení motorových vozidel, ve znění pozdějších předpisů \nv termínu:""")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    paragraf = doc.add_paragraph()
    run = paragraf.add_run(f' {datum.strftime("\n%d.%m.%Y")}')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = True
    run.underline = True

    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Evidenční číslo'
    hdr_cells[1].text = 'Jméno'
    hdr_cells[2].text = 'Přijmení'
    hdr_cells[3].text = 'Datum narození'
    hdr_cells[4].text = 'Skupina ŘO'
    hdr_cells[5].text = 'Druh zkoušky'
    hdr_cells[6].text = 'Začátek'
    for item in seznam_studentu: # cykl pro výpis studentů do tabulky
        row_cells = table.add_row().cells
        row_cells[0].text = item.zak.ev_cislo
        row_cells[1].text = item.zak.jmeno

        prijmeni_paragraph = row_cells[2].paragraphs[0]
        prijmeni_run = prijmeni_paragraph.add_run(item.zak.prijmeni)
        prijmeni_run.font.name = 'Arial'
        prijmeni_run.font.size = Pt(10)
        prijmeni_run.bold = True  # Nastavení tučného písma

        row_cells[3].text = item.zak.narozeni.strftime("%d.%m.%Y")
        row_cells[4].text = item.typ_zkousky
        row_cells[5].text = item.druh_zkousky
        row_cells[6].text = item.zacatek.strftime("%H:%M")

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                paragraph.paragraph_format.line_spacing = Pt(15)  # Nastavení výšky řádku (např. 15 pt)
                paragraph.paragraph_format.space_after = Pt(5)    # Mezera pod odstavcem
                paragraph.paragraph_format.space_before = Pt(5)   # Mezera nad odstavcem

    paragraf = doc.add_paragraph()
    run = paragraf.add_run("\nZkouška výše uvedených uchazečů se bude konat v ")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run = paragraf.add_run("sídle Magistrátu města Mostu, v ulici Radniční 1/2, IV. patro.")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = True
    paragraf = doc.add_paragraph()
    paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = paragraf.add_run("\nPotřebné doklady (občanský průkaz popř. potvrzení dlouhodobého pobytu, žádost uchazeče, lékařskou prohlídku, průkaz žáka, třídní knihu a potvrzení o zaplacení správního poplatku) předložte, prosím, ke kontrole zkušebnímu komisaři týž den před zkouškou, nebo po dohodě s komisařem i dříve.")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    paragraf = doc.add_paragraph()
    run = paragraf.add_run("\nS pozdravem")
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    #Footer
    footer = doc.sections[0].footer

    table = footer.add_table(rows=1, cols=2, width=Pt(500))
    table.autofit = False

    table.columns[0].width = Pt(250)
    table.columns[1].width = Pt(250)

    left_cell = table.cell(0, 0)
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    left_paragraph.add_run("Magistrát města Mostu").bold = True
    left_paragraph.add_run("\nRadniční 1/2, 434 01 Most")
    left_paragraph.add_run("\nTel.: +420 476 448 111")
    left_paragraph.add_run("\nposta@mesto-most.cz")
    left_paragraph.add_run("\nwww.mesto-most.cz")
    for run in left_paragraph.runs:
        run.font.size = Pt(8)
        run.font.name = 'Work Sans'

    right_cell = table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    right_paragraph.add_run("IČO: 00266094")
    right_paragraph.add_run("\nDIČ: CZ00266094")
    right_paragraph.add_run("\nID datové schránky: pftb6uv")
    right_paragraph.add_run("\nČ. ú.: 1041386359/0800")
    for run in right_paragraph.runs:
        run.font.size = Pt(8)
        run.font.name = 'Work Sans'
    
    paragraf = footer.add_paragraph()
    run = paragraf.add_run("Toto je automaticky generovaný dokument.")
    run.font.size = Pt(8)
    run.font.name = 'Work Sans'

    return doc

def create_document_end(studenti, autoskola):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Nadpis 1
    nadpis = document.add_paragraph(f"{autoskola.nazev}")
    nadpis.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nadpis.runs[0].bold = True
    nadpis.runs[0].font.size = Pt(20)

    # Nadpis 2
    nadpis = document.add_paragraph("SEZNAM ŽADATELŮ")
    nadpis.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nadpis.runs[0].bold = True
    nadpis.runs[0].font.size = Pt(14)

    document.add_paragraph()  # prázdný řádek

    # Hlavička tabulky
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    header = table.rows[0].cells
    hlavicky = ["Evidenční číslo", "Jméno Příjmení", "Datum narození", "Adresa", "Číslo ŘP", "Skupina vozidel", "Zahájení/\nukončení kurzu"]
    for i, text in enumerate(hlavicky):
        p = header[i].paragraphs[0]
        p.add_run(text).bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data studentů
    for student in studenti:
        row = table.add_row().cells
        cele_jmeno = f"{student['jmeno']} {student['prijmeni']}"
        datum_narozeni = student['narozeni'].strftime("\n%d.%m.%Y")
        datum_zacatek = student['zacatek'].strftime("%d. %m. %Y")
        datum_konec = student['konec'].strftime("%d. %m. %Y")

        row[0].text = student['ev_cislo']
        row[1].text = cele_jmeno
        row[2].text = datum_narozeni
        row[3].text = student['adresa']
        row[4].text = student['cislo'] if student['cislo'] else '----------'
        row[5].text = student['skupina']             # Skupina vozidel
        row[6].text = f"{datum_zacatek}\n{datum_konec}"

    slozka = "Ukonceni_vycviku"
    os.makedirs(slozka, exist_ok=True)

    # Název souboru = nazev_autoškoly_datum_konce.docx
    nazev_autoskoly = autoskola.nazev.replace(" ", "_")  
    datum_posledniho = max(s['konec'] for s in studenti)
    datum_str = datum_posledniho.strftime("%Y-%m-%d")
    nazev_souboru = f"{nazev_autoskoly}_{datum_str}.docx"
    # Celá cesta
    cesta = os.path.join(slozka, nazev_souboru)
    document.save(cesta)