from flask import Flask, render_template, jsonify, request, redirect, url_for, flash, abort, send_file, session, send_from_directory, current_app #mikrorámec na routování a celkovou správu webu
from flask_login import LoginManager, login_user, logout_user, login_required, current_user #podpůrná knihovna na správu přihlášených uživatelů
from flask_mail import Mail, Message #podpůrná knihovna na posílání emailů
from database import db, Zak, Termin, Autoskola, Zaznam, Komisar, Zapsany_zak, Vozidlo, Upozorneni, Superadmin #ORM modely na komunikaci s databází
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature #knihovna na generování tokenů. Kontrola, jestli není token 'prošlí' a je autentický 'neupravený'.
from sqlalchemy import or_, desc #metoda na možnost or v quary
from sqlalchemy.sql import func
from docx import Document #Objekt, který generuje word dokument z kódu
import app_logic #soubor s metodamy
import os
from werkzeug.utils import safe_join
from urllib.parse import unquote
from config_copy import Config #nastavení pro posílání emailů
from io import BytesIO 
from hashlib import sha256  #hashovací metoda
from datetime import date, datetime, timedelta
from pathlib import Path

app = Flask(__name__) # Vytvoření instance pro web
app.config['SECRET_KEY'] = Config.SECRET_KEY

# Nastavení aplikace pro spoj s databází
app.config['SQLALCHEMY_DATABASE_URI'] = Config.DB_MODEL_PATH
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SESSION_PERMANENT'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=60) # nastavuje jak dlouho může být uživatel neaktivní před tím než se ukončí session

app.config.from_object(Config) # Nastavení na připojení na email server
db.init_app(app) # Spojení s databází
loginManager = LoginManager(app) # Instance spravuje přihlášené uživatele
mail = Mail(app) # Instance na posílaní mailů
loginManager.login_view = 'home'
serializer = URLSafeTimedSerializer(app.config['SECRET_KEY']) # instance objektu, který vytváří token na resetování a vytváření hesla

@app.route("/login", methods=['GET', 'POST'])
@app.route("/home", methods=['GET', 'POST'])
@app.route("/", methods=['GET', 'POST'])
def home():
    """
    Stránka slouží k přihlášení uživatelů.

    Funkcionalita:
        - Zpracovává GET a POST requesty.
        - Při GET requestu vrací přihlašovací formulář.
        - Při POST requestu ověřuje přihlašovací údaje uživatele.

    Args:
        email (str): Emailová adresa zadaná uživatelem při POST requestu.
        heslo (str): Heslo zadané uživatelem při POST requestu.

    Returns:
        - render_template("home.html"): HTML stránka s přihlašovacím formulářem (GET request).
        - redirect(url_for('calendar')): Přesměrování na kalendář komisaře (POST request, správné přihlášení).
        - redirect(url_for('calendar')): Přesměrování na kalendář autoškoly (POST request, správné přihlášení).
        - flash("message.html", message=error): Chybová zpráva při špatném přihlášení.
    """ 
    if current_user.is_authenticated:
        return redirect(url_for('calendar'))
    if request.method == 'POST': # při POST requestu
        email = request.form['login_email'] # načte si email
        heslo = request.form['login_heslo'] # načte si string z formu pro heslo
        if email.endswith('@mesto-most.cz'): # kontrolo jestli email končí @mesto-most.cz pro admina
            superadmin = Superadmin.query.filter_by(email=email).first() # v DB se pokusíme najít superadmina, jestli email nepatří jemu
            if superadmin:
                if app_logic.porovnat_hesla(heslo, superadmin.heslo): # tady se porovná heslo z databáze a hashovaná forma zadaného hesla
                    login_user(app_logic.User(superadmin.id)) # přihlásí uživatele a vytvoří mu session
                    return redirect(url_for('calendar'))
            cil = Komisar.query.filter_by(email=email).first()  # podle emailu se najde komisař
            if cil:
                if app_logic.porovnat_hesla(heslo, cil.heslo): 
                    if cil.isAdmin: # Podmínka zjišťuje, jestli má sloupec isAdmin hodnotu 
                        session['isAdmin'] = True # do sessionu se přidá údaj o tom, že příhlášený mam admin práva
                        login_user(app_logic.User(cil.id, True)) # Pokud ano, tak přidá komisaři isAdmin = True
                    else:
                        login_user(app_logic.User(cil.id)) # Pokud ne, tak přidá komisaři isAdmin = False
                    return redirect(url_for('calendar'))
                else:
                    flash('Neplatný email nebo heslo', category='mess_error')
            else:
                flash('Neplatný email nebo heslo', category='mess_error')
        else: # zbytek ifu jestki email končí @mesto-most.cz
            cil = Autoskola.query.filter_by(email=email).first()
            if cil:
                if app_logic.porovnat_hesla(heslo, cil.heslo):
                    login_user(app_logic.User(cil.id))
                    zaznam = Zaznam(druh='přihlásil se', kdy=datetime.now(), zprava='Autoškola se přihlásila do aplikace', id_autoskoly=current_user.id) # vytvoří záznam o tom, že se autoškola přihlásila
                    db.session.add(zaznam) # tady se záznam přidá do db.session
                    db.session.commit() # tady se db.session nahraje do databáze
                    return redirect(url_for('calendar'))
                else:
                    flash('Neplatný email nebo heslo', category='mess_error')
            else:
                flash('Neplatný email nebo heslo', category='mess_error')

    return render_template('home.html')

@app.route('/zapomenute_heslo', methods=['GET', 'POST'])
def forgotten_password():
    """
    Zpracovává funkci zapomenutého hesla, včetně ověření emailu, generování tokenu a odesílání odkazů na resetování hesla.

    Tato metoda zpracovává GET a POST požadavky pro stránku "zapomenuté heslo". 
    Autentizovaní uživatelé jsou přesměrováni na stránku kalendáře. 
    Neautentizovaní uživatelé mohou zadat emailovou adresu, která se ověří, 
    identifikuje uživatele (administrátor nebo autoškola), vytvoří resetovací token a odešle email s odkazem na resetování hesla.

    Vrací:
        - Přesměruje autentizované uživatele na stránku kalendáře.
        - Při GET požadavku zobrazí HTML stránku pro zapomenuté heslo.
        - Při úspěšném POST požadavku odešle email s odkazem na resetování hesla nebo zobrazí odpovídající chybovou hlášku.
        - Po zpracování POST požadavku přesměruje na domovskou stránku.

    Vyvolává:
        - Exception: Pokud se nepodaří odeslat email, zobrazí se uživateli hláška a bude přesměrován na domovskou stránku.

    Metody:
        - GET: Zobrazí HTML šablonu pro zapomenuté heslo.
        - POST: Zpracuje požadavek na odeslání emailu pro resetování hesla.
    """
    if current_user.is_authenticated:
        return redirect(url_for('calendar'))
    
    if request.method == 'GET':
        return render_template('zapomenute_heslo.html')
    
    if request.method == 'POST':
        email = request.form['email']
        if email.endswith('@mesto-most.cz'): # kontrolo jestli email nekončí @mesto-most.cz pro admina
            cil = Komisar.query.filter_by(email=email).first() # pokud jo, tak se ho pokusí najít
        else:
            cil = Autoskola.query.filter_by(email=email).first() # pokud ne, tak ho hledá v autoškolách
        
        if cil: # pokud to najde komisaře nebo i autoškolu
            token = serializer.dumps(email, salt=Config.SALT_1) # vytvoří token z emailu a solí
            reset_url = url_for('reset_password', token=token, _external=True) # vytvoří odkaz na resetování hesla

            msg = Message(
                subject="Resetování hesla",
                sender='ReTesty@mesto-most.cz',
                recipients=[str(email)],
                body=f'Klikněte na tento odkaz pro resetování hesla: {reset_url}. Odkaz bude aktivní po dobu dvou hodin.',
                html=f'<p>Klikněte na tento odkaz pro resetování hesla:</p><a href="{reset_url}">Resetovat heslo</a>'
            )    
            try:
                mail.send(msg)
                flash('Email s odkazem na resetování hesla byl poslán.', category='success')
                return redirect(url_for('home'))
            except Exception as e:
                flash(f'Email se nepodařilo odeslat. Prosím, kontaktujte Magistrát města Most.', category='error')
                return redirect(url_for('home'))

    return redirect(url_for('home'))

@app.route('/resetovat/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """
    Zpracovává resetování hesla na základě platného tokenu.

    Tento endpoint obsluhuje GET a POST požadavky pro resetování hesla:
    - V GET požadavku zobrazí formulář pro nastavení nového hesla.
    - V POST požadavku ověří token, aktualizuje heslo v databázi a přesměruje uživatele na domovskou stránku.

    Vrací:
        - Při GET požadavku šablonu `reset_hesla.html`.
        - Při úspěšném POST požadavku přesměruje uživatele na domovskou stránku.
        - Pokud je token neplatný nebo vypršel, zobrazí příslušnou chybovou zprávu.

    Vyvolává:
        - `SignatureExpired`: Pokud vyprší platnost tokenu (po 2 hodinách).
        - Vrací zprávu "Nestihnul jsi to".
        - `BadSignature`: Pokud je token změněn nebo neplatný.
        - Vyvolá chybu 404.

    Metody:
        - GET:
            - Ověří platnost tokenu.
            - Zobrazí šablonu `reset_hesla.html` pro zadání nového hesla.
        - POST:
            - Ověří platnost tokenu.
            - Aktualizuje heslo v databázi podle typu uživatele:
                - Pokud email končí na `@mesto-most.cz`, heslo je změněno u komisaře.
                - Jinak je heslo změněno u autoškoly.
            - Používá SHA-256 hash pro ukládání hesel.
            - Ukládá změny do databáze a přesměruje uživatele na domovskou stránku.
    """
    try:
        email = serializer.loads(token, salt=Config.SALT_1, max_age=7200) # token se rozbalí, pokud bude "prošlí" vyhodí SignatureExpired, pokud bude token změněn vyhodí BadSignature
        if request.method == 'GET': # Pokud je token v pořádku a request je GET
            return render_template('reset_hesla.html')
        
        if request.method == 'POST': 
            nove_heslo= request.form['password'] # heslo z formuláře při POST          
            if email.endswith('@mesto-most.cz'): # kontrola jak končí email aby se hledalo ve správné tabulce
                cil = Komisar.query.filter_by(email=email).first()
                if not cil:
                    cil = Superadmin.query.filter_by(email=email).first()
                cil.heslo = sha256(nove_heslo.encode('utf-8')).hexdigest()
                db.session.commit()
                return redirect(url_for('home'))
            else:
                autoskola = Autoskola.query.filter_by(email=email).first()
                autoskola.heslo = sha256(nove_heslo.encode('utf-8')).hexdigest()
                db.session.commit()
                return redirect(url_for('home'))
            
    except SignatureExpired: # Pokud vyprší platnost tokenu
        return 'Token vypršel'
    except BadSignature: # Token je neplatný nebo byl změněn (např. útok nebo modifikace)
        abort(404)

@app.route('/nove_heslo/<token>', methods=['GET', 'POST'])
def create_password(token):
    """
    Zpracovává vytvoření nového hesla na základě platného tokenu.

    Endpoint umožňuje resetování hesla pro uživatele pomocí tokenu zaslaného emailem.
    Token je platný po dobu dvou dnů (172800 sekund). Podporuje GET a POST požadavky:
    - V GET požadavku zobrazí formulář pro zadání nového hesla.
    - V POST požadavku uloží nové heslo do databáze.

    Args:
        token (str): Token generovaný při procesu žádosti o resetování hesla, obsahuje zakódovaný email uživatele.

    Vrací:
        - Při GET požadavku šablonu `reset_hesla.html` pro zadání nového hesla.
        - Při úspěšném POST požadavku přesměruje uživatele na domovskou stránku.
        - Pokud je token neplatný nebo vypršel, zobrazí příslušnou chybovou zprávu nebo vyvolá chybu 404.

    Vyvolává:
        - `SignatureExpired`: Pokud vyprší platnost tokenu (po 2 dnech).
        - Vrací zprávu "Nestihnul jsi to".
        - `BadSignature`: Pokud je token změněn nebo neplatný.
        - Vyvolá chybu 404.

    Metody:
        - GET:
            - Ověří platnost tokenu.
            - Zobrazí šablonu `reset_hesla.html` pro zadání nového hesla.
        - POST:
            - Ověří platnost tokenu.
            - Z formuláře načte nové heslo.
            - Heslo zahashuje pomocí SHA-256 a uloží do databáze pro uživatele (autoškolu) odpovídající emailu z tokenu.
            - Uloží změny do databáze a přesměruje uživatele na domovskou stránku.
    """
    try:
        email = serializer.loads(token, salt=Config.SALT_2, max_age=172800) # token je platný dva dny
        
        if request.method == 'POST':
            nove_heslo= request.form['password']
            autoskola = Autoskola.query.filter_by(email=email).first()
            autoskola.heslo = sha256(nove_heslo.encode('utf-8')).hexdigest()
            db.session.commit()
            return redirect(url_for('home'))
        else:
            return render_template('reset_hesla.html') #vrací sice html na reset hesla ale funkčnost je správná

    except SignatureExpired: # Pokud vyprší platnost tokenu
        return 'Nestihnul jsi to'
    except BadSignature: # Token je neplatný nebo byl změněn (např. útok nebo modifikace)
        abort(404)    

@app.route('/calendar', methods=['GET'])
@login_required
def calendar():
    """
    Zpracovává zobrazení hlavní stránky kalendáře podle role aktuálně přihlášeného uživatele.

    Tento endpoint kontroluje roli přihlášeného uživatele a vrací odpovídající verzi hlavní stránky kalendáře.

    Vrací:
        - Šablonu `main_page.html` s parametry podle role uživatele:
            - `superadmin=True`: Pokud je uživatel superadministrátor.
            - `admin=True`: Pokud je uživatel komisař s administrátorskými právy.
            - Bez parametrů: Pokud je uživatel běžný komisař.
        - Šablonu `main_page_user.html`: Pokud je uživatel běžný uživatel (např. autoškola).

    Vyvolává:
        - `404`: Pokud uživatel není přihlášený (zajištěno pomocí `@login_required`).

    Metody:
        - GET:
            - Kontroluje roli uživatele:
                - **SuperAdmin**: Vrací stránku `main_page.html` s příznakem `superadmin=True`.
                - **Admin Komisař**: Vrací stránku `main_page.html` s příznakem `admin=True`.
                - **Komisař bez práv**: Vrací stránku `main_page.html` bez dalších příznaků.
                - **Běžný uživatel (autoškola)**: Vrací stránku `main_page_user.html`.

    Předpoklady:
        - Uživatel musí být přihlášený.
        - Objekt `current_user` obsahuje atributy:
            - `isSuperAdmin` (bool): Určuje, zda je uživatel superadministrátor.
            - `isCommissar` (bool): Určuje, zda je uživatel komisař.
            - `isAdmin` (bool): Určuje, zda je uživatel administrátor.
    """

    if current_user.isSuperAdmin:
        return render_template('main_page.html', superadmin=True)
    elif current_user.isCommissar:
        if current_user.isAdmin:
            return render_template('main_page.html', admin=True)
        else:
            return render_template('main_page.html')
    else:
        return render_template('main_page_user.html')

@app.route('/term/<id>', methods=['GET'])
@login_required
def term(id):
    """
    Zpracovává zobrazení a správu termínu na základě role aktuálně přihlášeného uživatele a stavu termínu.

    Tento endpoint umožňuje uživatelům a administrátorům přístup k informacím o termínech, zápis studentů, 
    správu zkoušek a nastavení termínů. Obsah stránky a dostupné akce se liší podle role uživatele (komisař, admin, superadmin) 
    a aktuálního stavu termínu (aktivní, neaktivní, uzavřený).

    Args:
        id (str): Identifikátor termínu, který se má načíst.

    Vrací:
        - `404`: Pokud termín neexistuje nebo je neaktivní pro uživatele.
        - Šablony HTML podle role uživatele a stavu termínu:
            - `term.html`: Stránka pro uživatele s možností zapsat žáky nebo zobrazit jejich stav.
            - `term_admin.html`: Stránka pro adminy a superadminy s kompletní správou termínu.
            - `term_komisar.html`: Stránka pro komisaře bez možnosti správy studentů.
            - `zapis_studenta_adminem.html`: Stránka pro adminy pro zápis studentů do neaktivního termínu.
            - `term_conclusion.html`: Stránka se závěry termínu pro adminy a superadminy.

    Vyvolává:
        - `404`: Pokud termín neexistuje nebo není dostupný uživateli.
        - Různé výjimky z databáze v případě chybného dotazu.

    Chování podle role:
        - **Uživatel**:
            - Má přístup pouze k termínům, které jsou aktivní ('Y') a uzavřené ('R').
            - Může zapsat studenty nebo zobrazit seznam zapsaných studentů.
        - **Komisař**:
            - Vidí termíny a seznamy studentů podle autoškol.
        - **Admin/Superadmin**:
            - Vidí všechny termíny.
            - Může přidávat studenty, nastavovat začátek zkoušky a potvrzovat studenty.
            - Má přístup ke správě neaktivních termínů a uzavřených termínů.

    Stavy termínu:
        - **'N' (Neaktivní)**:
            - Umožňuje adminům připravit termín k aktivaci, přidat studenty a nastavit detaily.
        - **'Y' (Aktivní)**:
            - Uživatelé a admini mohou pracovat s aktivními záznamy studentů.
        - **'R' (Uzavřený)**:
            - Termín je uzavřen, admini mohou zobrazit závěry studentů a výsledky zkoušek.
    """

    termin = Termin.query.filter_by(id=id).first()

    if not termin: # podmínka zkontroluje jestli termín existuje
        abort(404)
    else:
        session['term_id'] = id #pokud termín existuje dáme ho do sessionu protože s ním budeme ještě pracovat

    if not current_user.isCommissar: 
        
        match termin.ac_flag:
            case 'N': # pokud je termín neaktivní abort 
                abort(404)
            case 'Y': # pokud je aktivní
                zaci = Zapsany_zak.query \
                    .join(Zak) \
                    .join(Termin) \
                    .filter(Zapsany_zak.id_terminu == id, Zak.id_autoskoly == current_user.id, or_(Zapsany_zak.potvrzeni == 'Y', Zapsany_zak.potvrzeni == 'W')) \
                    .all()
                
                zaci_Y = []
                zaci_W = []
                for item in zaci:
                    if item.potvrzeni == 'Y':
                        zaci_Y.append({
                            'id': item.zak.id,
                            'typ_zkousky': item.typ_zkousky,
                            'druh_zkousky': item.druh_zkousky,
                            'ev_cislo': item.zak.ev_cislo,
                            'jmeno': item.zak.jmeno,
                            'prijmeni': item.zak.prijmeni,
                            'narozeni': item.zak.narozeni,
                            'zacatek': item.zak.zacatek,
                            'konec': item.zak.konec,
                            'prvni': item.zak.prvni,
                            'potvrzeni': item.potvrzeni,
                            'cas': item.zacatek
                        })
                    if item.potvrzeni == 'W':
                        zaci_W.append({
                            'id': item.zak.id,
                            'typ_zkousky': item.typ_zkousky,
                            'druh_zkousky': item.druh_zkousky,
                            'ev_cislo': item.zak.ev_cislo,
                            'jmeno': item.zak.jmeno,
                            'prijmeni': item.zak.prijmeni,
                            'narozeni': item.zak.narozeni,
                            'zacatek': item.zak.zacatek,
                            'konec': item.zak.konec,
                            'prvni': item.zak.prvni,
                            'potvrzeni': item.potvrzeni
                        })

                volna_mista = termin.max_ridicu - Zapsany_zak.query.filter_by(id_terminu=termin.id, potvrzeni='Y').count() - len(zaci_W)

                return render_template('term.html', termin=termin, volna_mista=volna_mista,zaci_Y=zaci_Y, zaci_W=zaci_W)
            case 'R':
                zaci = Zapsany_zak.query \
                    .join(Zak) \
                    .join(Termin) \
                    .filter(Zapsany_zak.id_terminu == id, Zak.id_autoskoly == current_user.id, Zapsany_zak.potvrzeni == 'Y') \
                    .all()
                
                zaci_s = []
                for item in zaci:
                        zaci_s.append({
                            'id': item.zak.id,
                            'typ_zkousky': item.typ_zkousky,
                            'druh_zkousky': item.druh_zkousky,
                            'ev_cislo': item.zak.ev_cislo,
                            'jmeno': item.zak.jmeno,
                            'prijmeni': item.zak.prijmeni,
                            'narozeni': item.zak.narozeni,
                            'zacatek': item.zak.zacatek,
                            'konec': item.zak.konec,
                            'prvni': item.zak.prvni,
                            'potvrzeni': item.potvrzeni,
                            'zaver': item.zaver,
                            'cas': item.zacatek
                        })
                return render_template('term_read.html', termin=termin, zaci=zaci_s)
                 
    elif current_user.isCommissar:
        volna_mista = termin.max_ridicu - Zapsany_zak.query.filter_by(id_terminu=termin.id, potvrzeni='Y').count()
        match termin.ac_flag:
            case 'Y':
            # Vrátí všechny zapsané studenty, schromáždí je pod jejich autoškoly a zobrazí jestli jsou už přijmutí nebo ne!
                zaci = Zapsany_zak.query \
                    .join(Zak) \
                    .join(Termin) \
                    .filter(Zapsany_zak.id_terminu == id) \
                    .all()
                
                komisari = Komisar.query.all() #vrací list objektů Komisar
                autoskoly = Autoskola.query.all() #vrací list objektů Autoškola, který se v html přidá do selectu

                autoskoly_lst = []

                for autoskola in autoskoly:
                    autoskoly_lst.append({
                        'id': autoskola.id,
                        'nazev': autoskola.nazev
                    })

                zaci_v_as = {} # zde se jako klíče budou dávat název autoškol a hodnota bude list žáků
                for item in zaci:
                    autoskola = Autoskola.query.filter_by(id=item.zak.id_autoskoly).first()
                    komisar = next((k for k in komisari if k.id == item.id_komisare), None) 
                    if autoskola.nazev not in zaci_v_as:
                        zaci_v_as[autoskola.nazev] = [{     
                                                            'id': item.zak.id,                 
                                                            'ev_cislo': item.zak.ev_cislo,
                                                            'jmeno': item.zak.jmeno,
                                                            'prijmeni': item.zak.prijmeni,
                                                            'narozeni': item.zak.narozeni,
                                                            'typ_zkousky': item.typ_zkousky,
                                                            'druh_zkousky': item.druh_zkousky,
                                                            'zacatek': item.zak.zacatek,
                                                            'konec': item.zak.konec,
                                                            'prvni': item.zak.prvni,
                                                            'potvrzeni': item.potvrzeni,
                                                            'komisar': f'{komisar.jmeno} {komisar.prijmeni}' if komisar else None,
                                                            'cas': item.zacatek
                                                        }]
                    elif autoskola.nazev in zaci_v_as:
                        zaci_v_as[autoskola.nazev].append({
                                                            'id': item.zak.id,                 
                                                            'ev_cislo': item.zak.ev_cislo,
                                                            'jmeno': item.zak.jmeno,
                                                            'prijmeni': item.zak.prijmeni,
                                                            'narozeni': item.zak.narozeni,
                                                            'typ_zkousky': item.typ_zkousky,
                                                            'druh_zkousky': item.druh_zkousky,
                                                            'zacatek': item.zak.zacatek,
                                                            'konec': item.zak.konec,
                                                            'prvni': item.zak.prvni,
                                                            'potvrzeni': item.potvrzeni,
                                                            'komisar': f'{komisar.jmeno} {komisar.prijmeni}' if komisar else None,
                                                            'cas': item.zacatek
                                                        })

                srovnany_dict = dict(sorted(zaci_v_as.items()))
                print(srovnany_dict)
                if current_user.isAdmin or current_user.isSuperAdmin: # Vrací html, ve kterém lze zapsat studenty a jde vidět kdo, koho zkouší
                    return render_template('term_admin.html', superadmin=current_user.isSuperAdmin, admin= current_user.isAdmin, list_as=srovnany_dict, termin=termin, komisari= komisari, volna_mista= volna_mista, autoskoly=autoskoly_lst)
                else: # Vrací html, ve kterém nelze zapsat studenty a nejde vidět kdo, koho zkouší
                    return render_template('term_komisar.html', list_as=srovnany_dict, termin=termin, komisari= komisari, volna_mista= volna_mista, autoskoly=autoskoly_lst)

            case 'N':
                """
                V případě, že je termín stále neaktivní, admin by měl mít možnost od něj zapsat studenty sám bez
                zásahu autoškoly. Tudíž v tomto případě je potřeba dostat žáky, kteří jsou již zapsaní, čekají na zápis
                a formuláře pro zapsaní žáků.
                """
                # Vrátí všechny zapsané studenty, schromáždí je pod jejich autoškoly a zobrazí jestli jsou už přijmutí nebo ne!
                zaci = Zapsany_zak.query \
                    .join(Zak) \
                    .join(Termin) \
                    .filter(Zapsany_zak.id_terminu == id) \
                    .all()
                
                komisari = Komisar.query.all() #vrací list objektů Komisar
                autoskoly = Autoskola.query.all() #vrací list objektů Autoškola, který se v html přidá do selectu

                autoskoly_lst = []

                for autoskola in autoskoly:
                    autoskoly_lst.append({
                        'id': autoskola.id,
                        'nazev': autoskola.nazev
                    })

                zaci_v_as = {} # zde se jako klíče budou dávat název autoškol a hodnota bude list žáků
                for item in zaci:
                    autoskola = Autoskola.query.filter_by(id=item.zak.id_autoskoly).first()
                    komisar = next((k for k in komisari if k.id == item.id_komisare), None) 
                    if autoskola.nazev not in zaci_v_as:
                        zaci_v_as[autoskola.nazev] = [{     
                                                            'id': item.zak.id,                 
                                                            'ev_cislo': item.zak.ev_cislo,
                                                            'jmeno': item.zak.jmeno,
                                                            'prijmeni': item.zak.prijmeni,
                                                            'narozeni': item.zak.narozeni,
                                                            'typ_zkousky': item.typ_zkousky,
                                                            'druh_zkousky': item.druh_zkousky,
                                                            'zacatek': item.zak.zacatek,
                                                            'konec': item.zak.konec,
                                                            'prvni': item.zak.prvni,
                                                            'potvrzeni': item.potvrzeni,
                                                            'komisar': f'{komisar.jmeno} {komisar.prijmeni}' if komisar else None,
                                                            'cas': item.zacatek
                                                        }]
                    elif autoskola.nazev in zaci_v_as:
                         zaci_v_as[autoskola.nazev].append({
                                                            'id': item.zak.id,                 
                                                            'ev_cislo': item.zak.ev_cislo,
                                                            'jmeno': item.zak.jmeno,
                                                            'prijmeni': item.zak.prijmeni,
                                                            'narozeni': item.zak.narozeni,
                                                            'typ_zkousky': item.typ_zkousky,
                                                            'druh_zkousky': item.druh_zkousky,
                                                            'zacatek': item.zak.zacatek,
                                                            'konec': item.zak.konec,
                                                            'prvni': item.zak.prvni,
                                                            'potvrzeni': item.potvrzeni,
                                                            'komisar': f'{komisar.jmeno} {komisar.prijmeni}' if komisar else None,
                                                            'cas': item.zacatek
                                                        })

                srovnany_dict = dict(sorted(zaci_v_as.items()))
                volna_mista = termin.max_ridicu - len([zak for zak in zaci if zak.potvrzeni == 'Y'])
                return render_template('zapis_studenta_adminem.html', superadmin=current_user.isSuperAdmin, admin= current_user.isAdmin, termin=termin, list_as=srovnany_dict, komisari=komisari, volna_mista=volna_mista, autoskoly=autoskoly_lst)
            case 'R':
                zaci = Zapsany_zak.query \
                    .join(Zak) \
                    .join(Termin) \
                    .filter(Zapsany_zak.id_terminu == id) \
                    .all()
                
                komisari = Komisar.query.all() #vrací list objektů Komisar

                komisari_lst = []

                for komisar in komisari:
                    komisari_lst.append({
                        'id': komisar.id,
                        'jmeno': komisar.jmeno,
                        'prijmeni': komisar.prijmeni
                    })


                zaci_v_as = {} # zde se jako klíče budou dávat název autoškol a hodnota bude list žáků
                for item in zaci:
                    autoskola = Autoskola.query.filter_by(id=item.zak.id_autoskoly).first()
                    komisar = next((k for k in komisari if k.id == item.id_komisare), None) 
                    if autoskola.nazev not in zaci_v_as:
                        zaci_v_as[autoskola.nazev] = [{     
                                                            'id': item.zak.id,                 
                                                            'ev_cislo': item.zak.ev_cislo,
                                                            'jmeno': item.zak.jmeno,
                                                            'prijmeni': item.zak.prijmeni,
                                                            'narozeni': item.zak.narozeni,
                                                            'typ_zkousky': item.typ_zkousky,
                                                            'druh_zkousky': item.druh_zkousky,
                                                            'zacatek': item.zak.zacatek,
                                                            'konec': item.zak.konec,
                                                            'prvni': item.zak.prvni,
                                                            'potvrzeni': item.potvrzeni,
                                                            'komisar': f'{komisar.jmeno} {komisar.prijmeni}' if komisar else None,
                                                            'cas': item.zacatek,
                                                            'zaver': item.zaver
                                                        }]
                    elif autoskola.nazev in zaci_v_as:
                        zaci_v_as[autoskola.nazev].append({
                                                            'id': item.zak.id,                 
                                                            'ev_cislo': item.zak.ev_cislo,
                                                            'jmeno': item.zak.jmeno,
                                                            'prijmeni': item.zak.prijmeni,
                                                            'narozeni': item.zak.narozeni,
                                                            'typ_zkousky': item.typ_zkousky,
                                                            'druh_zkousky': item.druh_zkousky,
                                                            'zacatek': item.zak.zacatek,
                                                            'konec': item.zak.konec,
                                                            'prvni': item.zak.prvni,
                                                            'potvrzeni': item.potvrzeni,
                                                            'komisar': f'{komisar.jmeno} {komisar.prijmeni}' if komisar else None,
                                                            'cas': item.zacatek,
                                                            'zaver': item.zaver
                                                        })                   
                srovnany_dict = dict(sorted(zaci_v_as.items()))
                autoskoly = Autoskola.query.all()

                autoskoly_lst = []

                for autoskola in autoskoly:
                    autoskoly_lst.append({
                        'id': autoskola.id,
                        'nazev': autoskola.nazev
                    })
                return render_template('term_conclusion.html', superadmin = current_user.isSuperAdmin, admin = current_user.isAdmin, list_as=srovnany_dict, termin=termin, komisari= komisari_lst, autoskoly= autoskoly_lst)

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    """
    Zpracovává zobrazení profilu aktuálně přihlášeného uživatele.

    Endpoint se přizpůsobuje podle role uživatele:
    - **Autoškola**: Zobrazí detailní informace o autoškole, seznam vozidel a žáků.
    - **Komisař/Admin**: Umožňuje výběr autoškoly k zobrazení profilu.

    Vrací:
        - `profile.html`: Pokud je uživatel autoškola, zobrazí informace o ní.
        - `profile_admin.html`: Pokud je uživatel komisař/admin, zobrazí přehled autoškol.
        - Přesměrování: Pokud je odeslán POST request k výběru autoškoly.

    Vyvolává:
        - `flash()`: Chybová zpráva při nesprávném výběru autoškoly.
    
    Metody:
        - GET: Zobrazí stránku s informacemi o uživateli.
        - POST: Zpracuje výběr autoškoly administrátorem.
    """
    if current_user.isCommissar:
        if request.method == 'GET':
            autoskoly = Autoskola.query.all()

            return render_template('profile_admin.html', superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin, autoskoly=autoskoly)
        if request.method == 'POST':
            autoskola = request.form["driving-school"]
            if not autoskola:
                flash('Musíte vybrat autoškolu!', category='error')
            else:
                return redirect(url_for('profile_admin', id=autoskola))
    else:
        if request.method == 'GET':
            id = current_user.id
            skola = Autoskola.query.filter_by(id=id).first()
            seznam_vozidel =  Vozidlo.query.filter_by(id_autoskoly= current_user.id).all()
            studenti = (
                db.session.query(
                    Zak,
                    func.coalesce(Termin.datum, None).label("datum")  # Nahrazení NULL hodnoty
                )
                .outerjoin(Zapsany_zak, (Zak.id == Zapsany_zak.id_zaka) & (Zapsany_zak.zaver == 'W'))
                .outerjoin(Termin, Termin.id == Zapsany_zak.id_terminu)
                .filter(Zak.id_autoskoly == current_user.id)
                .all()
            )
            seznam_studentu = []

            for student, datum in studenti:
                seznam_studentu.append({
                    'id': student.id,
                    'ev_cislo': student.ev_cislo,
                    'jmeno': student.jmeno,
                    'prijmeni': student.prijmeni,
                    'zacatek': student.zacatek.strftime("%d.%m.%Y") if student.zacatek else '',
                    'konec': student.konec.strftime("%d.%m.%Y") if student.konec else '',
                    'termin': (
                        "splnil" if student.splnil  # Pokud student splnil, zobrazí "splnil"
                        else datum.strftime('%d.%m.%Y') if datum  # Pokud existuje termín, zobrazí datum
                        else "nepřihlášený"  # Jinak zobrazí "nepřihlášený"
                    )
                })

            return render_template('profile.html', vozidla= seznam_vozidel, autoskola=skola, studenti=seznam_studentu)
        if request.method == 'POST':
            pass

@app.route('/profile/<id>')
@login_required
def profile_admin(id):
    """
    Zpracovává zobrazení profilu autoškoly pro administrátory a komisaře.

    Endpoint umožňuje zobrazit podrobnosti o konkrétní autoškole, jejích vozidlech a žácích.

    Args:
        id (str): Identifikátor autoškoly.

    Vrací:
        - `profile_commissar.html`: Stránku s informacemi o autoškole, jejích vozidlech a žácích.
        - `404`: Pokud uživatel nemá oprávnění nebo autoškola neexistuje.

    Vyvolává:
        - `404`: Pokud uživatel není administrátor nebo komisař.
        - `Exception`: Pokud nastane chyba při získávání dat.

    Metody:
        - GET: Načte a zobrazí detaily autoškoly.
    """
    if not current_user.isCommissar:
        abort(404)
    try:
        autoskola = Autoskola.query.filter_by(id=id).first()
        vozidla = Vozidlo.query.filter_by(id_autoskoly=id).all()
        studenti = (
                db.session.query(
                    Zak,
                    func.coalesce(Termin.datum, None).label("datum")  # Nahrazení NULL hodnoty
                )
                .outerjoin(Zapsany_zak, (Zak.id == Zapsany_zak.id_zaka) & (Zapsany_zak.zaver == 'W'))
                .outerjoin(Termin, Termin.id == Zapsany_zak.id_terminu)
                .filter(Zak.id_autoskoly == id)
                .all()
            )

        autoskola_dic = {
            'id': autoskola.id,
            'nazev': autoskola.nazev,
            'email': autoskola.email,
            'adresa': autoskola.adresa_u,
            'datovka': autoskola.da_schranka
        }

        vozidla_lst = []
        zaci_lst = []

        for vozidlo in vozidla:
            vozidla_lst.append({
                'id': vozidlo.id,
                'znacka': vozidlo.znacka,
                'model': vozidlo.model,
                'rz': vozidlo.spz
            })
        for student, datum in studenti:
            zaci_lst.append({
                'id': student.id,
                    'ev': student.ev_cislo,
                    'jmeno': student.jmeno,
                    'prijmeni': student.prijmeni,
                    'zacatek': student.zacatek.strftime("%d.%m.%Y") if student.zacatek else '',
                    'konec': student.konec.strftime("%d.%m.%Y") if student.konec else '',
                    'termin': (
                        "splnil" if student.splnil  # Pokud student splnil, zobrazí "splnil"
                        else datum.strftime('%d.%m.%Y') if datum  # Pokud existuje termín, zobrazí datum
                        else "nepřihlášený"  # Jinak zobrazí "nepřihlášený"
                    )
            })

        return render_template('profile_commissar.html', superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin, autoskola=autoskola_dic, vozidla=vozidla_lst, zaci=zaci_lst)

    except Exception as e:
        string = "" 
        if not autoskola:
             string += " není autoškola,"
        elif not vozidla:
            string += " nejsou vozidla,"
        elif not studenti:
            string += " nejsou zaci,"
        else:
            string += " Něco je špatně jinde"
        return string

@app.route('/new_driving_school', methods=['GET', 'POST'])
@login_required
def new_driving_school():
    """
    Zpracovává přidávání nové autoškoly do systému.

    Endpoint umožňuje administrátorům přidávat nové autoškoly. Podporuje GET a POST požadavky:
    - V GET požadavku zobrazí formulář pro přidání nové autoškoly.
    - V POST požadavku ověří, zda autoškola již existuje. Pokud ne, přidá ji do databáze, 
    vygeneruje token a odešle email s odkazem na vytvoření hesla.

    Vrací:
        - `404`: Pokud uživatel není administrátor.
        - Šablona `new_driving_school.html` při GET požadavku.
        - Přesměrování na domovskou stránku při úspěšném vytvoření autoškoly.
        - Přesměrování na stránku kalendáře při chybě v odeslání emailu.

    Vyvolává:
        - `404`: Pokud uživatel není administrátor.
        - Výjimky při selhání odesílání emailu.

    Metody:
        - GET:
            - Zobrazí formulář pro přidání nové autoškoly.
            - Zahrnuje informace o roli uživatele (admin/superadmin).
        - POST:
            - Zpracuje údaje z formuláře:
                - `nazev` (str): Název autoškoly.
                - `dat_schranka` (str): Datová schránka autoškoly.
                - `email` (str): Email autoškoly.
            - Ověří, zda autoškola již existuje:
                - Pokud neexistuje:
                    - Přidá autoškolu do databáze.
                    - Vygeneruje token a odešle email s odkazem na vytvoření hesla.
                - Pokud existuje:
                    - Vygeneruje nový token a odešle email s odkazem na vytvoření hesla.
            - Zobrazuje flash zprávy o stavu operace:
                - Úspěšné přidání autoškoly.
                - Chyba při odesílání emailu.

    Předpoklady:
        - Uživatel musí být přihlášený a mít roli administrátora.
        - Objekt `current_user` obsahuje atributy:
            - `isAdmin` (bool): Určuje, zda je uživatel administrátor.
            - `isSuperAdmin` (bool): Určuje, zda je uživatel super administrátor.

    """
    if not current_user.isAdmin:
        abort(404)
    if request.method == 'POST':
        nazev = request.form.get('nazev')
        dat_schranka = request.form.get('datova-schranka')
        email =  request.form.get('email')
        adresa = request.form.get('adresa')
        
        autoskola = Autoskola.query.filter_by(email=email).first()

        if not autoskola:
            nova_autoskola = Autoskola(nazev=nazev, da_schranka=dat_schranka, email=email, adresa_u=adresa)
            db.session.add(nova_autoskola)
            db.session.commit()

            flash('Nová autoškola přidána', category='success')

            token = serializer.dumps(email, salt=Config.SALT_2) # vytvoří token z emailu a soli
            newPassword_url = url_for('create_password', token=token, _external=True) # vytvoří odkaz na resetování hesla

            msg = Message(
                subject="Odkaz na vytvoření hesla",
                sender='ReTesty@mesto-most.cz',
                recipients=[str(email)],
                body=f'Kliknutím na tento odkaz bude přesměrováni na stránku pro vytvoření hesla: {newPassword_url}. Odkaz bude aktivní po dobu dvou dní.',
                html=f'<p>Klikněte na tento odkaz pro vytvoření hesla:</p><a href="{newPassword_url}">Vytvořit heslo</a><p>Odkaz bude aktivní po dobu dvou dní.</p>'
                #TODO přidat soubor s manuálem
            )    
            try:
                mail.send(msg)
                flash('Email s odkazem byl poslán.', category='success')
                return redirect(url_for('calendar'))
            except Exception as e:
                flash(f'Email se nepodařilo odeslat. Prosím, kontaktujte Magistrát města Most.', category='error')
                return redirect(url_for('calendar'))
        
        if autoskola:
            token = serializer.dumps(email, salt=Config.SALT_2) # vytvoří token z emailu a soli
            newPassword_url = url_for('create_password', token=token, _external=True) # vytvoří odkaz na resetování hesla

            msg = Message(
                subject="Odkaz na vytvoření hesla",
                sender='ReTesty@mesto-most.cz',
                recipients=[str(email)],
                body=f'Klikněte na tento odkaz bude přesměrováni na stránku pro vytvoření hesla: {newPassword_url}. Odkaz bude aktivní po dobu dvou dní.',
                html=f'<p>Klikněte na tento odkaz pro vytvoření hesla:</p><a href="{newPassword_url}">Vytvořit heslo</a><p>Odkaz bude aktivní po dobu dvou dní.</p>'
                #TODO přidat soubor s manuálem
            )    
            try:
                mail.send(msg)
                flash('Email s odkazem byl poslán.', category='success')
                return redirect(url_for('home'))
            except Exception as e:
                flash(f'Email se nepodařilo odeslat. Prosím, kontaktujte Magistrát města Most.', category='error')
                return redirect(url_for('calendar'))    


        flash('Nová autoškola přidána', category='success')
        return render_template('new_driving_school.html')
    
    if request.method == 'GET':
        return render_template('new_driving_school.html', superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin)

@app.route('/teaching_training', methods=['GET'])
@login_required
def teaching_training():
    """
    Zpracovává žádosti o zápis do výcviku a vrací příslušnou stránku podle role přihlášeného uživatele.

    Endpoint poskytuje odlišné šablony a data podle role uživatele:
    - **Komisař**: Získá seznam všech autoškol a zobrazí stránku pro zápis adminem.
    - **Autoškola**: Získá seznam vozidel své autoškoly a zobrazí stránku pro zápis žáků.

    Vrací:
        - Šablonu `sign_up_admin.html` s údaji o autoškolách, pokud je uživatel komisař.
        - Šablonu `sign_up.html` s údaji o autoškole a vozidlech, pokud je uživatel autoškola.
        - `404`: Pokud požadavek neodpovídá žádné z podmínek.

    Vyvolává:
        - `404`: Pokud požadavek neodpovídá žádné z podmínek.

    Metody:
        - GET:
            - Zobrazuje příslušnou stránku podle role uživatele.

    Chování podle role:
        - **Komisař**:
            - Načte všechny autoškoly z databáze.
            - Vrací šablonu `sign_up_admin.html` s parametry:
                - `autoskoly` (list): Seznam autoškol (ID a název).
                - `superadmin` (bool): Indikuje, zda je uživatel superadmin.
                - `admin` (bool): Indikuje, zda je uživatel admin.
        - **Autoškola**:
            - Načte detailní informace o aktuální autoškole a její vozidla z databáze.
            - Vrací šablonu `sign_up.html` s parametry:
                - `vozidla` (list): Seznam vozidel autoškoly.
                - `autoskola` (object): Detailní informace o autoškole.

    Předpoklady:
        - Uživatel musí být přihlášený.
        - Objekt `current_user` obsahuje atributy:
            - `isCommissar` (bool): Určuje, zda je uživatel komisař.
            - `id` (int): Identifikátor přihlášené autoškoly.
            - `isAdmin` (bool): Určuje, zda je uživatel admin.
            - `isSuperAdmin` (bool): Určuje, zda je uživatel superadmin.
    """
    if current_user.isCommissar:
        ausk = Autoskola.query.all() # z db si vytahneme všechny autoškoly

        autoskoly = []
        for skola in ausk:
            autoskoly.append({
                'id': skola.id,
                'nazev': skola.nazev,
                'adresa': skola.adresa_u
            })

        return render_template('sign_up_admin.html', autoskoly=autoskoly, superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin)
    else:
        autoskola = Autoskola.query.filter_by(id=current_user.id).first()
        vozidla = Vozidlo.query.filter_by(id_autoskoly=current_user.id).all()
        return render_template('sign_up.html', vozidla=vozidla, autoskola=autoskola)
    
    abort(404)

@app.route('/logs', methods= ['GET', 'POST'])
@login_required
def logs():
    """
    Zpracovává zobrazení logů autoškoly, přístupné pouze administrátorům.

    Tento endpoint umožňuje administrátorům zobrazit logy autoškol na základě zadaných filtrů:
    - Administrátor si vybere autoškolu, typ záznamu a/nebo datum zobrazení.
    - Endpoint vyhledá odpovídající logy v databázi a zobrazí je.

    Vrací:
        - `404`: Pokud uživatel není administrátor.
        - Šablonu `logs.html`:
            - Při GET požadavku zobrazí formulář s výběrem autoškoly a filtrů.
            - Při POST požadavku zobrazí logy na základě zadaných filtrů.

    Vyvolává:
        - `404`: Pokud uživatel není administrátor.

    Metody:
        - GET:
            - Vrací formulář s možností výběru:
                - Autoškoly.
                - Typu záznamu.
                - Datum od.
            - Parametry předané šabloně:
                - `autoskoly` (list): Seznam autoškol s jejich ID a názvy.
        - POST:
            - Načte údaje z formuláře:
                - `autoskola` (int): ID autoškoly (volitelné).
                - `druh` (str): Typ logu (volitelné).
                - `datum` (str): Datum od (volitelné).
            - Filtruje záznamy v databázi na základě zadaných parametrů:
                - Pokud je zadána autoškola, vyhledá logy pouze pro ni.
                - Pokud je zadán typ logu, vyhledá odpovídající záznamy.
                - Pokud je zadáno datum, vyhledá záznamy od tohoto data.
            - Vrací šablonu `logs.html` s parametry:
                - `logs` (list): Seznam nalezených logů (datum, typ, zpráva, název autoškoly).
                - `autoskoly` (list): Seznam autoškol s jejich ID a názvy.
                - `superadmin` (bool): Určuje, zda je uživatel superadmin.
                - `admin` (bool): Určuje, zda je uživatel admin.

    Předpoklady:
        - Uživatel musí být přihlášený a mít roli administrátora.
        - Objekt `current_user` obsahuje atributy:
            - `isAdmin` (bool): Určuje, zda je uživatel admin.
            - `isSuperAdmin` (bool): Určuje, zda je uživatel superadmin.
        - Objekt `Zaznam` obsahuje:
            - `kdy` (datetime): Datum a čas záznamu.
            - `druh` (str): Typ záznamu.
            - `zprava` (str): Obsah záznamu.
            - `id_autoskoly` (int): ID autoškoly, které se záznam týká.
            - `autoskola`: Relace na objekt autoškoly (pro název).

    Omezení:
        - Vrací maximálně 100 záznamů na základě filtrů.
    """

    if not current_user.isAdmin:
        abort(404)

    autoskoly = Autoskola.query.all()
        # Pokud přibudou typy záznamu tak je můžeme zjistit takhle: druhy_zaznamu = [choice for choice in Zaznam.druh.type.enums]
    lst_as = []
    for autoskola in autoskoly:
        lst_as.append({'id':autoskola.id,
                        'nazev':autoskola.nazev})

    if request.method == 'GET':
        return render_template('logs.html', superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin, autoskoly=lst_as)
    
    if request.method == 'POST':
        id_as = request.form.get('autoskola') or None
        druh = request.form.get('druh') or None
        datum = request.form.get('datum') or None

        quary = Zaznam.query # Tady se přípravý quary, které potom budu doplňovat podle parametrů

        if id_as:
            quary = quary.filter(Zaznam.id_autoskoly == id_as)
        if druh:
            quary = quary.filter(Zaznam.druh == druh)
        if datum:
            quary = quary.filter(Zaznam.kdy >= datum)

        quary.limit(100) # limit na počet záznamů
        zaznamy = quary.order_by(desc(Zaznam.kdy)).all()

        lst_zaznamy= []
        for zaznam in zaznamy:
            lst_zaznamy.append({
                'kdy': zaznam.kdy,
                'druh': zaznam.druh,
                'zprava': zaznam.zprava,
                'autoskola': zaznam.autoskola.nazev
            })
            
        return render_template('logs.html', superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin, logs=lst_zaznamy, autoskoly=lst_as)

@app.route('/rights', methods= ['GET'])
@login_required
def rozdat_prava():
    """
    Zpracovává správu práv pro komisaře, přístupné pouze superadministrátorům.

    Endpoint umožňuje superadministrátorům zobrazit seznam všech komisařů a jejich aktuální stav práv. 
    Slouží k zobrazení a případné správě práv (aktuální implementace pouze zobrazuje data).

    Vrací:
        - `404`: Pokud uživatel není superadministrátor.
        - Šablonu `rights.html` s přehledem všech komisařů a jejich práv.

    Vyvolává:
        - `404`: Pokud uživatel není superadministrátor.

    Metody:
        - GET:
            - Načte seznam všech komisařů z databáze.
            - Zformátuje jejich data do seznamu obsahujícího:
                - `id` (int): ID komisaře.
                - `jmeno` (str): Jméno komisaře.
                - `prijmeni` (str): Příjmení komisaře.
                - `isAdmin` (str): Datum, kdy byla komisaři přidělena administrátorská práva (ve formátu "DD.MM.RRRR"), 
                nebo `None`, pokud práva nemá.
            - Předá zformátovaný seznam šabloně `rights.html`.

    Předpoklady:
        - Uživatel musí být přihlášený a mít roli superadministrátora.
        - Objekt `current_user` obsahuje atribut:
            - `isSuperAdmin` (bool): Určuje, zda je uživatel superadministrátor.
        - Objekt `Komisar` obsahuje:
            - `id` (int): ID komisaře.
            - `jmeno` (str): Jméno komisaře.
            - `prijmeni` (str): Příjmení komisaře.
            - `isAdmin` (datetime or None): Datum přiřazení administrátorských práv, nebo `None`.
    """

    if not current_user.isSuperAdmin:
        abort(404)
    
    komisari = Komisar.query.all()

    lst_komisari = []
    for komisar in komisari:
        lst_komisari.append({
            'id': komisar.id,
            'jmeno': komisar.jmeno,
            'prijmeni': komisar.prijmeni,
            'isAdmin': date.today() < komisar.end_isAdmin if komisar.end_isAdmin else False,
            'start_isAdmin': komisar.start_isAdmin.strftime("%d.%m.%Y") if komisar.start_isAdmin else None,
            'end_isAdmin': komisar.end_isAdmin.strftime("%d.%m.%Y") if komisar.end_isAdmin else None
        })

    return render_template('rights.html', komisari=lst_komisari)

@login_required
@app.route("/end_of_training_site")
def end_of_training_view():
    if not current_user.isCommissar:
        abort(404)

    slozka = "Ukonceni_vycviku"
    nazvy_souboru = []

    if os.path.exists(slozka):
        for filename in os.listdir(slozka):
            if filename.endswith(".docx"):
                cesta = os.path.join(slozka, filename)
                cas_zmeny = os.path.getmtime(cesta)  # poslední změna v sekundách (timestamp)
                nazvy_souboru.append((cas_zmeny, filename))

    # Seřadíme sestupně (nejnovější nahoře)
    nazvy_souboru.sort(reverse=True)

    # Předáme jen názvy (bez časů)
    return render_template("end_of_training_site.html", soubory=[f[1] for f in nazvy_souboru], superadmin=current_user.isSuperAdmin, admin=current_user.isAdmin,)

#API metody
@login_required
@app.route('/api/get_document/<path:name>', methods=['GET'])
def get_end_document(name):
    name = unquote(name)

    slozka_absolutni = os.path.abspath(os.path.join(current_app.root_path, "..", "Ukonceni_vycviku"))
    cesta = safe_join(slozka_absolutni, name)

    print("Absolutní složka:", slozka_absolutni)
    print("Celá cesta k souboru:", cesta)

    if not os.path.isfile(cesta):
        abort(404)

    return send_from_directory(directory=slozka_absolutni, path=name, as_attachment=True)

@login_required
@app.route('/pridat_prava', methods= ['POST'])
def pridat_prava():
    """
    Uděluje administrátorská práva komisaři na základě ID a data.

    Endpoint umožňuje přiřadit komisaři administrátorská práva do určitého data. 
    Vloží zadané datum do sloupce `isAdmin` v databázové tabulce komisařů. 
    Po uplynutí tohoto data systém automaticky odstraní hodnotu ze sloupce `isAdmin`, 
    čímž dojde ke zrušení administrátorských práv.

    Vrací:
        - `200`: Pokud byly práva úspěšně přiřazeny.
        - `404`: Pokud nebyl nalezen komisař s daným ID nebo pokud se na endpoint snaží dostat někdo jiný než superadmin.
        - `400`: Pokud nastane jakákoli jiná chyba při zpracování požadavku.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování požadavku.

    Metody:
        - POST:
            - Přijímá JSON objekt s klíči:
                - `id` (int): ID komisaře, kterému mají být udělena práva.
                - `date` (str): Datum ve formátu `RRRR-MM-DD`, do kterého mají být práva platná.
            - Načte komisaře z databáze podle zadaného ID.
            - Pokud je komisař nalezen:
                - Aktualizuje jeho sloupec `isAdmin` na zadané datum.
                - Uloží změny do databáze.
                - Vrátí zprávu o úspěšném provedení.
            - Pokud komisař není nalezen:
                - Vrátí chybu `404` s příslušnou zprávou.
    """
    
    try:
        data = request.get_json()
        id = data.get('id')
        date_start = data.get('date_start')
        date_end = data.get('date_end')    

        komisar = Komisar.query.filter_by(id=id).first()
        if komisar:
            komisar.start_isAdmin = date_start
            komisar.end_isAdmin = date_end
            db.session.commit()
            return jsonify({"message": "Data přijata úspěšně"}), 200 # Odeslání odpovědi o úspěchu
        else:
            return jsonify({"error": "Komisař nebyl nalezen"}), 404
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400

@login_required
@app.route('/odebrat_prava', methods= ['POST'])
def odebrat_prava():
    """
    Odebírá administrátorská práva komisaři na základě ID.

    Endpoint umožňuje odstranit administrátorská práva komisaře. 
    Nastaví hodnotu sloupce `isAdmin` na `None` v databázové tabulce komisařů.

    Vrací:
        - `200`: Pokud byla práva úspěšně odebrána.
        - `404`: Pokud nebyl nalezen komisař s daným ID.
        - `400`: Pokud nastane jakákoli jiná chyba při zpracování požadavku.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování požadavku.

    Metody:
        - POST:
            - Přijímá JSON objekt s klíčem:
                - `id` (int): ID komisaře, kterému mají být odebrána práva.
            - Načte komisaře z databáze podle zadaného ID.
            - Pokud je komisař nalezen:
                - Nastaví jeho sloupec `isAdmin` na `None`.
                - Uloží změny do databáze.
                - Vrátí zprávu o úspěšném odebrání práv.
            - Pokud komisař není nalezen:
                - Vrátí chybu `404` s příslušnou zprávou.
    """
    if not current_user.isSuperAdmin:
        abort(404)
    try:
        data = request.get_json()
        id = data.get("id")  # Získání id z JSON objektu
        
        # Najít komisaře podle id
        komisar = Komisar.query.filter_by(id=id).first()
        
        if komisar:
            komisar.isAdmin = False
            komisar.start_isAdmin = None
            komisar.end_isAdmin = None
            db.session.commit()
            return jsonify({"message": "Práva byla úspěšně odebrána"}), 200
        else:
            return jsonify({"error": "Komisař nebyl nalezen"}), 404
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400

@app.route('/create_autoskola', methods=['POST'])
def nova_autoskola():
    """
    Vytváří nový záznam autoškoly v databázi.

    Endpoint umožňuje přidat novou autoškolu pomocí POST požadavku. 
    Data autoškoly (název, datová schránka, adresa učebny, email a heslo) 
    jsou přijímána z formuláře. Heslo je před uložením zahashováno pomocí SHA256.

    Vrací:
        - Přesměrování na endpoint `admin` s potvrzením o úspěšném vytvoření autoškoly.
        - Vyvolá chybu (Exception), pokud dojde k selhání při ukládání do databáze.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při vytváření nebo ukládání autoškoly.

    Metody:
        - POST:
            - Přijímá následující formulářová data:
                - `nazev` (str): Název autoškoly.
                - `datova_schranka` (str): Datová schránka autoškoly.
                - `adresa_ucebny` (str): Adresa učebny autoškoly.
                - `email` (str): E-mailová adresa autoškoly.
                - `heslo` (str): Heslo autoškoly (zahashované před uložením).
            - Vytvoří nový záznam autoškoly v databázi.
            - Uloží nový záznam pomocí `db.session.commit()`.
            - Při úspěšném vytvoření zobrazí uživateli flash zprávu `Autoškola se přidala!
    """
    nazev = request.form['nazev']
    da_schranka = request.form['datova_schranka']
    adresa_u = request.form['adresa_ucebny']
    email = request.form['email']
    heslo = sha256(request.form['heslo'].encode('utf-8')).hexdigest() # heslo se zahashuje
    try:
        autoskola = Autoskola(nazev=nazev, da_schranka=da_schranka, adresa_u=adresa_u, email=email, heslo=heslo)
        db.session.add(autoskola)
        db.session.commit()
    except:
        raise Exception()
    else:
        flash('Autoškola se přidala!', category='mess_success')
        return redirect(url_for('admin'))

@app.route('/create_komisar', methods=['POST'])
def novy_komisar():
    """
    API metoda pro vytvoření komisaře, určená pro použití pouze během testování

    Parametry:
        request: myslím si, že se jedná o objekt, který v sobě přenáší JSON z formu pro informace o komisaři

    Vrací:
        raise (str): pokud při vytvoření a nebo commitu do databáze vznikne error
        flash(str): při úspěšném přidání se vytvoří zpráva, která se zobrazí při dalším render_templatu
        redirect(url_for(str)): po přijmutí a vytvoření komisaře vrátí uživatele zpátky na /admin
    """
    email = request.form['email']
    heslo = sha256(request.form['heslo'].encode('utf-8')).hexdigest()
    jmeno = request.form['jmeno']
    prijmeni = request.form['prijmeni']
    try:
        komisar = Komisar(email=email, heslo=heslo, jmeno=jmeno, prijmeni=prijmeni)
        db.session.add(komisar)
        db.session.commit()
    except:
        raise 'Někde je chyba'
    else:
        flash('Komisař se přidal!', category='mess_success')
        return redirect(url_for('admin'))

@app.route('/create_zak', methods=['POST'])
def novy_zak():
    """
    API metoda pro vytvoření žáka, určená pro použití pouze během testování

    Parametry:
        request: myslím si, že se jedná o objekt, který v sobě přenáší JSON z formu pro informace o žákovi

    Vrací:
        str: pokud při vytvoření a nebo commitu do databáze vznikne error
        flash(str): při úspěšném přidání se vytvoří zpráva, která se zobrazí při dalším render_templatu
        redirect(url_for(str)): po přijmutí a vytvoření autoškoly vrátí uživatele zpátky na /admin
    """
    ev_cislo = request.form['ev_cislo']
    jmeno = request.form['jmeno']
    prijmeni = request.form['prijmeni']
    narozeni = request.form['narozeni']
    adresa = request.form['adresa']
    id_autoskoly = request.form['id_autoskoly']

    try:
        zak = Zak(ev_cislo=ev_cislo, jmeno=jmeno, prijmeni=prijmeni, narozeni=narozeni, adresa=adresa, id_autoskoly=id_autoskoly)
        db.session.add(zak)
        db.session.commit()
    except:
        raise 'Někde je chyba'
    else:
        flash('Žák se přidal!', category='mess_success')
        return redirect(url_for('admin'))    

@app.route('/create_termin', methods=['POST'])
def novy_termin():
    """
    API metoda pro vytvoření termínu, určená pro použití pouze během testování

    Parametry:
        request: myslím si, že se jedná o objekt, který v sobě přenáší JSON z formu pro informace o termínu

    Vrací:
        str: pokud při vytvoření a nebo commitu do databáze vznikne error
        flash(str): při úspěšném přidání se vytvoří zpráva, která se zobrazí při dalším render_templatu
        redirect(url_for(str)): po přijmutí a vytvoření autoškoly vrátí uživatele zpátky na /admin
    """
    if not (current_user.isCommissar or current_user.isAdmin):
       abort(404)

    datum = request.form['datum']
    max_ridicu = request.form['max_ridicu']

    try:
        termin = Termin(datum=datum, max_ridicu=max_ridicu)
        db.session.add(termin)
        db.session.commit()
    except:
        raise 'Někde je chyba'
    else:
        flash('Termín se přidal!', category='mess_success')
        return redirect(url_for('admin'))

@login_required
@app.route('/api/get_student', methods=['POST'])
def dostan_studenta():
    """
    API metoda sloužící pro nalezení žáka a kontrole jestli ještě není nikde jinde zapsaný. Metoda vrací žáka, aby ho připravila na zapis. 
    K metodě mají přístup pouze komisaři a admini.
    
    Parametry:
        request: Z FE obdržíme JSON s id autoškoly žáka, a Ev. Číslo žáka.
    
    Vrací:
        JSON 400: pokud kombinace id autoškoly a ev. čísla nikoho nenajde nebo pokud už je na termínu
        JSON 200: pokud vše proběhne v pořádku
        JSON 500: internal error
    """
    if not (current_user.isCommissar or current_user.isAdmin):
       abort(404)
    try:
        # Získání dat z požadavku
        data = request.get_json()
        autoskola_id = data.get('autoskolaId')
        evidencni_cislo = data.get('evidencniCislo')
        # Validace vstupů
        if not autoskola_id or not evidencni_cislo:
            return jsonify({"error": "Chybí id autoškoly nebo evidenční číslo"}), 400
        
        zak = Zak.query.filter_by(ev_cislo=evidencni_cislo, id_autoskoly=autoskola_id).order_by(desc(Zak.id)).first()

        if zak:
            if not zak.konec:
                return jsonify({"error": f"Žák nemá ukončený výcvik."}), 400
            zap_zak = Zapsany_zak.query.filter_by(id_zaka=zak.id, zaver='W').first()
            if zap_zak:
                return jsonify({"error": "Žák je již na zapsaný na termínu."}), 400
            zak_info = {
                'id': zak.id,
                'ev_cislo': zak.ev_cislo,
                'jmeno': zak.jmeno,
                'prijmeni': zak.prijmeni,
                'dat_nar': zak.narozeni,
                'adresa': zak.adresa,
                'zacatek': zak.zacatek,
                'konec': zak.konec if zak.konec else None,
                'prvni':zak.prvni if zak.prvni else None
                }
        else:
            return jsonify({"error": "Žák nebyl nalezen"}), 400
        return jsonify(zak_info), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@login_required
@app.route('/api/get_student_as', methods=['POST'])
def dostan_studenta_as():
    """
    Získá informace o žákovi na základě evidenčního čísla a ID autoškoly.

    Endpoint umožňuje získat informace o konkrétním žákovi v autoškole, pokud není již přihlášen na termín nebo nesplnil zkoušku. 
    Požadavek musí být proveden s platným JSON objektem obsahujícím evidenční číslo žáka.

    Vrací:
        - 200: Pokud je žák nalezen a nejsou žádná omezení, vrací JSON objekt s informacemi o žákovi.
        - 400: Pokud chybí vstupní data nebo pokud:
            - Žák již splnil zkoušku.
            - Žák je již přihlášen na termín.
            - Žák nebyl nalezen v databázi.
        - 404: Pokud je uživatel komisář (nemá oprávnění pro tento endpoint).
        - 500: Pokud dojde k interní chybě serveru.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování požadavku nebo přístupu k databázi.

    Metody:
        - POST:
            - Přijímá JSON objekt s klíči:
                - `evidencniCislo` (str): Evidenční číslo žáka.
            - Ověří, zda je uživatel oprávněn požadavek provést.
            - Ověří, zda žák existuje a splňuje podmínky pro získání informací.
            - Vrací informace o žákovi jako JSON objekt nebo chybovou odpověď.
    """

    try:
        # Získání dat z požadavku
        if current_user.isCommissar:
            abort(404)
        data = request.get_json()
        autoskola_id = current_user.id
        evidencni_cislo = data.get('evidencniCislo')

        # Validace vstupů
        if not autoskola_id or not evidencni_cislo:
            return jsonify({"error": "Chybí id autoškoly nebo evidenční číslo"}), 400
        
        zak = Zak.query.filter_by(ev_cislo=evidencni_cislo, id_autoskoly=autoskola_id).order_by(desc(Zak.id)).first()

    
        if zak:
            if not zak.konec:
                return jsonify({"error": f"Žák nemá ukončený výcvik."}), 400
            if zak.splnil:
                return jsonify({"error": "Žák je již splnil zkoušku."}), 400
            zap_zak = Zapsany_zak.query.filter_by(id_zaka=zak.id, zaver='W').first()
            if zap_zak:
                return jsonify({"error": "Žák je již zapsaný na termín."}), 400
            zak_info = {
                'id': zak.id,
                'ev_cislo': zak.ev_cislo,
                'jmeno': zak.jmeno,
                'prijmeni': zak.prijmeni,
                'dat_nar': zak.narozeni,
                'adresa': zak.adresa,
                'zacatek': zak.zacatek,
                'konec': zak.konec if zak.konec else '',
                'prvni':zak.prvni if zak.prvni else ''
                }
        else:
            return jsonify({"error": "Žák nebyl nalezen"}), 400
        
        return jsonify(zak_info), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@login_required
@app.route('/api/get_student_end', methods=['POST'])
def dostan_studenta_end():
    """
    API metoda sloužící pro nalezení žáka a kontrole jestli ještě není nikde jinde zapsaný. Metoda vrací žáka, aby ho připravila na zapis. 
    K metodě mají přístup pouze komisaři a admini.
    
    Parametry:
        request: Z FE obdržíme JSON s id autoškoly žáka, a Ev. Číslo žáka.
    
    Vrací:
        JSON 400: pokud kombinace id autoškoly a ev. čísla nikoho nenajde nebo pokud už je na termínu
        JSON 200: pokud vše proběhne v pořádku
        JSON 500: internal error
    """
    if not (current_user.isCommissar or current_user.isAdmin):
       abort(404)
    try:
        # Získání dat z požadavku
        data = request.get_json()
        autoskola_id = data.get('autoskolaId')
        evidencni_cislo = data.get('evidencniCislo')
        # Validace vstupů
        if not autoskola_id or not evidencni_cislo:
            return jsonify({"error": "Chybí id autoškoly nebo evidenční číslo"}), 400
        
        zak = Zak.query.filter_by(ev_cislo=evidencni_cislo, id_autoskoly=autoskola_id).order_by(desc(Zak.id)).first()

        if zak:
            zak_info = {
                'id': zak.id,
                'ev_cislo': zak.ev_cislo,
                'jmeno': zak.jmeno,
                'prijmeni': zak.prijmeni,
                'dat_nar': zak.narozeni,
                'adresa': zak.adresa,
                'zacatek': zak.zacatek,
                'konec': zak.konec if zak.konec else None,
                'prvni':zak.prvni if zak.prvni else None
                }
        else:
            return jsonify({"error": "Žák nebyl nalezen"}), 400
        return jsonify(zak_info), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@login_required
@app.route('/api/get_student_end_as', methods=['POST'])
def dostan_studenta_end_as():
    """
    API metoda sloužící pro nalezení žáka a kontrole jestli ještě není nikde jinde zapsaný. Metoda vrací žáka, aby ho připravila na zapis. 
    K metodě mají přístup pouze komisaři a admini.
    
    Parametry:
        request: Z FE obdržíme JSON s id autoškoly žáka, a Ev. Číslo žáka.
    
    Vrací:
        JSON 400: pokud kombinace id autoškoly a ev. čísla nikoho nenajde nebo pokud už je na termínu
        JSON 200: pokud vše proběhne v pořádku
        JSON 500: internal error
    """
    try:
        # Získání dat z požadavku
        data = request.get_json()
        evidencni_cislo = data.get('evidencniCislo')
        # Validace vstupů
        if not evidencni_cislo:
            return jsonify({"error": "Chybí evidenční číslo"}), 400
        
        zak = Zak.query.filter_by(ev_cislo=evidencni_cislo, id_autoskoly=current_user.id).order_by(desc(Zak.id)).first()

        if zak:
            zak_info = {
                'id': zak.id,
                'ev_cislo': zak.ev_cislo,
                'jmeno': zak.jmeno,
                'prijmeni': zak.prijmeni,
                'dat_nar': zak.narozeni,
                'adresa': zak.adresa,
                'zacatek': zak.zacatek,
                'konec': zak.konec if zak.konec else None,
                'prvni':zak.prvni if zak.prvni else None
                }
        else:
            return jsonify({"error": "Žák nebyl nalezen"}), 400
        return jsonify(zak_info), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@login_required
@app.route('/api/end_of_training_as', methods=['POST'])
def end_of_training_as():   
    try:
        data = request.get_json()
        studenti = data.get("studenti") # z dat dostane [{'id': int, 'datum':'yyyy-mm-dd'}, atd.]

        if not studenti or not isinstance(studenti, list):
            return jsonify({"error": "Neplatná nebo chybějící data o studentech."}), 400
        studenti_do_dokumentu = []
        aktualizovano = []
        for student in studenti:
            student_id = student.get("id")
            datum_konce = student.get("datum_konce")
            skupiny = student.get("skupiny_opravneni")

            if not student_id or not datum_konce:
                return jsonify({"error": f"Chybí id nebo datum konce u jednoho ze záznamů."}), 400

            # Ověření správného formátu datumu
            try:
                datum_obj = datetime.strptime(datum_konce, "%Y-%m-%d").date()
            except ValueError:
                return jsonify({"error": f"Neplatný formát datumu: {datum_konce}. Použijte YYYY-MM-DD."}), 400

            # Najít žáka
            zak = Zak.query.filter_by(id=student_id).first()
            if not zak:
                return jsonify({"error": f"Žák s ID {student_id} nebyl nalezen."}), 404
            if zak.konec:
                aktualizovano.append(f"Studentovi/tce {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl změněn datum konce výuky.")

            # Aktualizace datumu konce výcviku
            zak.konec = datum_obj
            studenti_do_dokumentu.append({
                'ev_cislo': zak.ev_cislo,
                'jmeno': zak.jmeno,
                'prijmeni': zak.prijmeni,
                'narozeni': zak.narozeni,
                'adresa':   zak.adresa,
                'zacatek': zak.zacatek,
                'konec': zak.konec,
                'skupina': skupiny,
                'cislo': zak.cislo if zak.cislo else None
            })
            
        autoskola = Autoskola.query.filter_by(id= current_user.id).first()
        db.session.commit()

        app_logic.create_document_end(studenti_do_dokumentu, autoskola)

        return jsonify({"success": True, "zaznamy": aktualizovano}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@login_required
@app.route('/api/end_of_training', methods=['POST'])
def end_of_training():
    if not (current_user.isCommissar or current_user.isAdmin):
       abort(404)   
    try:
        data = request.get_json()
        studenti = data.get("studenti") # z dat dostane [{'id': int, 'datum':'yyyy-mm-dd'}, atd.]

        if not studenti or not isinstance(studenti, list):
            return jsonify({"error": "Neplatná nebo chybějící data o studentech."}), 400

        autoskola_id = None
        studenti_do_dokumentu = []
        aktualizovano = []

        for student in studenti:
            student_id = student.get("id")
            datum_konce = student.get("datum_konce")
            skupiny = student.get("skupiny_opravneni")

            if not student_id or not datum_konce:
                return jsonify({"error": f"Chybí id nebo datum konce u jednoho ze záznamů."}), 400

            # Ověření správného formátu datumu
            try:
                datum_obj = datetime.strptime(datum_konce, "%Y-%m-%d").date()
            except ValueError:
                return jsonify({"error": f"Neplatný formát datumu: {datum_konce}. Použijte YYYY-MM-DD."}), 400
            # Najít žáka
            zak = Zak.query.filter_by(id=student_id).first()
            if not autoskola_id:
                autoskola_id= zak.id_autoskoly
            if not zak:
                return jsonify({"error": f"Žák s ID {student_id} nebyl nalezen."}), 404
            if zak.konec:
                aktualizovano.append(f"Studentovi/tce {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl změněn datum konce výuky.")
                upozorneni= Upozorneni(zprava= f'Studentovi/ce {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl změněn konec výcviku.',
                                    id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            # Aktualizace datumu konce výcviku
            else:
                upozorneni= Upozorneni(zprava= f'Studentovi/ce {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl přidán konec výcviku.',
                            id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            db.session.add(upozorneni)
            zak.konec = datum_obj
            studenti_do_dokumentu.append({
                'ev_cislo': zak.ev_cislo,
                'jmeno': zak.jmeno,
                'prijmeni': zak.prijmeni,
                'narozeni': zak.narozeni,
                'adresa':   zak.adresa,
                'zacatek': zak.zacatek,
                'konec': zak.konec,
                'skupina': skupiny,
                'cislo': zak.cislo if zak.cislo else None
            })
            
        autoskola = Autoskola.query.filter_by(id=autoskola_id).first()
        db.session.commit()

        app_logic.create_document_end(studenti_do_dokumentu, autoskola)

        return jsonify({"success": True, "zaznamy": aktualizovano}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@login_required
@app.route('/api/create_term', methods=['POST'])
def vytvor_termin():
    """
    Vytváří nový termín v systému na základě zadaných údajů.

    Endpoint umožňuje přidat nový termín pomocí POST požadavku. Data termínu (datum a maximální počet řidičů) 
    jsou přijímána jako JSON objekt. Při úspěšném vytvoření termínu dojde k přesměrování na stránku kalendáře.

    Vrací:
        - Přesměrování na endpoint `calendar` po úspěšném vytvoření termínu.
        - `400`: Pokud nastane jakákoli chyba při zpracování požadavku.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování požadavku nebo při ukládání do databáze.

    Metody:
        - POST:
            - Přijímá JSON objekt s klíči:
                - `dayId` (str): Datum termínu ve formátu `RRRR-MM-DD`.
                - `pocetMist` (int): Maximální počet řidičů, kteří se mohou přihlásit na termín.
            - Vytvoří nový záznam termínu v databázi pomocí modelu `Termin`.
            - Uloží nový záznam pomocí `db.session.commit()`.
            - Při úspěšném vytvoření přesměruje uživatele na endpoint `calendar`.
    """
    if not (current_user.isCommissar or current_user.isAdmin):
       abort(404)
    data = request.get_json()

    try:
        print(data.get('dayId'), data.get('pocetMist'))
        termin = Termin(datum=data.get('dayId'), max_ridicu=data.get('pocetMist'))
        db.session.add(termin)
        db.session.commit()
        return redirect(url_for('calendar'))
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400
    
@login_required
@app.route('/calendar_api', methods=['GET'])
def get_calendar_dates():
    """
    Vrací seznam termínů z kalendáře v závislosti na roli uživatele.

    Endpoint umožňuje získat seznam termínů v systému. Pro komisaře vrací všechny termíny, 
    zatímco pro ostatní uživatele (například autoškoly) vrací pouze aktivní termíny a ukončené termíny 
    spojené s jejich autoškolou.

    Vrací:
        - JSON seznam obsahující informace o termínech.
        - `400`: Pokud nastane jakákoli chyba při zpracování požadavku.

    Metody:
        - GET:
            - Pokud je uživatel komisař (`current_user.isCommissar == True`):
                - Načte všechny termíny z databáze.
                - Serializuje každý termín do formátu.
                - Vrátí seznam všech termínů ve formátu JSON.

            - Pokud uživatel není komisař (`current_user.isCommissar == False`):
                - Načte aktivní termíny (`ac_flag = 'Y'`) z databáze.
                - Načte ukončené termíny (`ac_flag = 'R'`) spojené s autoškolou uživatele (`current_user.id`).
                - Spojí aktivní i ukončené termíny dohromady.
                - Serializuje každý termín do formátu.
                - Vrátí seznam těchto termínů ve formátu JSON.
    """
    try:
        if current_user.isCommissar:
            terminy = Termin.query.all()
            
            # Serializace dat
            terminy_list = []
            for termin in terminy:
                termin_data = {
                    'id': termin.id,
                    'date': termin.datum.isoformat(),  # datum převedeme na string ve formátu ISO
                    'ac_flag': termin.ac_flag,
                    'max_ridicu': termin.max_ridicu,
                    'zapsani_zaci': len([z for z in termin.zapsani_zaci if z.potvrzeni == 'Y'])  # Počet zapsaných žáků na termín
                }
                terminy_list.append(termin_data)
            return jsonify(terminy_list)
        if not current_user.isCommissar:
            # Dotaz pro aktivní termíny (ac_flag = 'Y')
            aktivni_terminy = db.session.query(Termin).filter(Termin.ac_flag == 'Y').all()
            
            # Dotaz pro ukončené termíny (ac_flag = 'R') s žáky z konkrétní autoškoly
            ukoncene_terminy = db.session.query(Termin)\
                .join(Zapsany_zak)\
                .filter(Termin.ac_flag == 'R', Zapsany_zak.id_autoskoly == current_user.id, Zapsany_zak.potvrzeni == 'Y')\
                .all()

            # Spojíme aktivní i ukončené termíny dohromady
            terminy = aktivni_terminy + ukoncene_terminy

            # Příprava dat pro frontend ve formátu JSON
            terminy_data = []
            for termin in terminy:
                filtered_zapsani_zaci = [
                    zapsany for zapsany in termin.zapsani_zaci
                    if zapsany.id_autoskoly == current_user.id
                ]
            
                terminy_data.append({
                    "id": termin.id,
                    "date": termin.datum.isoformat(),
                    "ac_flag": termin.ac_flag,
                    "max_ridicu": termin.max_ridicu,
                    "zapsani_zaci": len(filtered_zapsani_zaci)
                })
            
            return jsonify(terminy_data)
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400 

@login_required
@app.route('/api/add_drivers', methods=['POST'])
def add_drivers():
    """
    API metoda pro zápis žáků na termín, určená pro komisaře.

    Endpoint umožňuje komisařům zapisovat žáky na konkrétní termín zkoušek. 
    Přijímá JSON objekt obsahující seznam žáků s detaily a vytváří záznamy v databázi 
    pro každého validního žáka.

    Vrací:
        - `200`: Pokud byl zápis úspěšný.
        - `400`: Pokud nastane jakákoli chyba při zpracování požadavku.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování nebo při ukládání do databáze.

    Metody:
        - POST:
            - Přijímá JSON seznam objektů s detaily žáků:
            - Pro každého žáka:
                - Zkontroluje existenci v databázi (`Zak`) na základě údajů.
                - Pokud žák existuje:
                    - Vytvoří záznam (`Zapsany_zak`) s údaji o termínu a autoškole.
                    - Vytvoří logovací záznam (`Zaznam`) o přidání žáka na termín.
                    - Uloží oba záznamy do databáze.
            - Vrátí zprávu o úspěchu (`200`), pokud byl zápis dokončen.
    """
    if not (current_user.isCommissar or current_user.isAdmin):
       abort(404)
    try:
        # Získání dat z požadavku
        data = request.get_json()
        students = data.get('students', [])

        # Validace vstupů
        if not students or not isinstance(students, list):
            return jsonify({"error": "Seznam studentů je prázdný nebo neplatný"}), 400
      
        # Získání termínu ze session
        term_id = session.get('term_id')
        if not term_id:
            return jsonify({"error": "ID termínu není k dispozici"}), 400
        
        zapsani_studenti = []
        for student in students:
            student_id = student.get('id')
            zak = Zak.query.filter_by(id= student_id).first()
            license_category = student.get('license_category')
            exam_type = student.get('exam_type')
            
            if not zak or not license_category or not exam_type:
                continue  # Přeskočit neplatné záznamy

            # Kontrola, zda student již není zapsán
            existing_entry = Zapsany_zak.query.filter_by(id_zaka=zak.id, id_terminu=term_id).first()
            if existing_entry:
                continue
            
            # Vytvoření nového záznamu
            new_entry = Zapsany_zak(
                id_zaka=zak.id,
                id_terminu=term_id,
                id_autoskoly=zak.id_autoskoly,
                typ_zkousky=license_category,
                druh_zkousky=exam_type.replace('_', ' ')
            )
            
            db.session.add(new_entry)
            zapsani_studenti.append(student)

        # Uložení změn do databáze
        db.session.commit()

        return jsonify({
            "success": True,
            "zapsani_studenti": zapsani_studenti,
            "pocet_zapsanych": len(zapsani_studenti),
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@login_required
@app.route('/api/add_drivers/r_term', methods=['POST'])
def add_drivers_r_term():
    try:
        data = request.get_json()
        students = data.get('students', [])

        # Validace vstupů
        if not students or not isinstance(students, list):
            return jsonify({"error": "Seznam studentů je prázdný nebo neplatný"}), 400
      
        # Získání termínu ze session
        term_id = session.get('term_id')
        if not term_id:
            return jsonify({"error": "ID termínu není k dispozici"}), 400
        
        termin = Termin.query.filter_by(id=term_id).first()
        zapsani_studenti = []
        for student in students:
            student_id = student.get('id')
            zak = Zak.query.filter_by(id= student_id).first()
            license_category = student.get('license_category')
            exam_type = student.get('exam_type')
            commisar = student.get('commisar')
            start = student.get('start_time')
            first = student.get('first_try')# kdy absolvoval první pokus
            
            if not zak or not license_category or not exam_type:
                continue  # Přeskočit neplatné záznamy

            # Kontrola, zda student již není zapsán
            existing_entry = Zapsany_zak.query.filter_by(id_zaka=zak.id, id_terminu=term_id).first()
            if existing_entry:
                continue
            
            # Vytvoření nového záznamu
            new_entry = Zapsany_zak(
                potvrzeni='Y',
                id_zaka=zak.id,
                id_terminu=term_id,
                id_autoskoly=zak.id_autoskoly,
                typ_zkousky=license_category,
                druh_zkousky=exam_type.replace('_', ' '),
                id_komisare = commisar,
                zacatek=start
                #prvni_pokus= first
            )
            upozorneni= Upozorneni(zprava= f'Student/ka {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl/a zapsán/a na termín {termin.datum.strftime("%d.%m.%Y")} v {start}',
                                   id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            db.session.add(upozorneni)
            db.session.add(new_entry)
            zapsani_studenti.append(student)

        # Uložení změn do databáze
        db.session.commit()

        return jsonify({
            "success": True,
            "zapsani_studenti": zapsani_studenti,
            "pocet_zapsanych": len(zapsani_studenti),
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@login_required
@app.route('/api/add_drivers_as', methods=['POST'])
def add_drivers_as():
    """
    API metoda pro zápis žáků na termín, určená pro autoškoly.

    Endpoint umožňuje autoškolám zapisovat své žáky na konkrétní termín zkoušek. 
    Přijímá JSON objekt obsahující seznam žáků s detaily a vytváří záznamy v databázi 
    pro každého validního žáka.

    Vrací:
        - `200`: Pokud byl zápis úspěšný.
        - `400`: Pokud nastane jakákoli chyba při zpracování požadavku.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování nebo při ukládání do databáze.

    Metody:
        - POST:
            - Přijímá JSON seznam objektů s detaily žáků:
            - Pro každého žáka:
                - Zkontroluje existenci v databázi (`Zak`) na základě údajů.
                - Pokud žák existuje:
                    - Vytvoří záznam (`Zapsany_zak`) s údaji o termínu a autoškole.
                    - Vytvoří logovací záznam (`Zaznam`) o přidání žáka na termín.
                    - Uloží oba záznamy do databáze.
            - Vrátí zprávu o úspěchu (`200`), pokud byl zápis dokončen.
    """
    try:
        # Získání dat z požadavku
        data = request.get_json()
        students = data.get('students', [])

        # Validace vstupů
        if not students or not isinstance(students, list):
            return jsonify({"error": "Seznam studentů je prázdný nebo neplatný"}), 400
      
        # Získání termínu ze session
        term_id = session.get('term_id')
        if not term_id:
            return jsonify({"error": "ID termínu není k dispozici"}), 400
        
        zapsani_studenti = []
        for student in students:
            student_id = student.get('id')
            zak = Zak.query.filter_by(id= student_id).first()
            license_category = student.get('license_category')
            exam_type = student.get('exam_type')
            if not zak.konec:
                return jsonify({'error': f'Zak {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} nemá ukončený výcvik.'})

            if not zak or not license_category or not exam_type:
                continue  # Přeskočit neplatné záznamy

            # Kontrola, zda student již není zapsán
            existing_entry = Zapsany_zak.query.filter_by(id_zaka=zak.id, id_terminu=term_id).first()
            if existing_entry:
                continue
            
            # Vytvoření nového záznamu
            new_entry = Zapsany_zak(
                id_zaka=zak.id,
                id_terminu=term_id,
                id_autoskoly=zak.id_autoskoly,
                typ_zkousky=license_category,
                druh_zkousky=exam_type.replace('_', ' ')
            )
            zaznam = Zaznam(druh='přidání', kdy=datetime.now(), zprava=f'Autoškola zapsala studentku/studenta {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} na termín s id:{term_id}.', id_autoskoly=autoskola.id)

            db.session.add(new_entry)
            zapsani_studenti.append(student)

        # Uložení změn do databáze
        db.session.commit()

        return jsonify({
            "success": True,
            "zapsani_studenti": zapsani_studenti,
            "pocet_zapsanych": len(zapsani_studenti),
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@login_required
@app.route('/api/enroll_by_admin', methods=['POST'])
def enroll_by_admin():
    """
    API metoda pro zápis žáků na termín správcem systému.

    Endpoint umožňuje administrátorovi zapsat žáky na termín zkoušek. 
    Přijímá JSON objekt obsahující seznam žáků s jejich detaily a pokouší se je najít v databázi.
    Pokud je žák nalezen, zapíše se na termín a vytvoří se záznam o zápisu. Pokud není nalezen, vrátí chybovou zprávu.

    Vrací:
        - `200`: Pokud byl zápis úspěšný pro všechny zadané žáky.
        - `400`: Pokud některý žák nebyl nalezen nebo nastala jiná chyba při zpracování požadavku.

    Vyvolává:
        - Výjimku (Exception): Pokud dojde k chybě při zpracování nebo při ukládání do databáze.

    Metody:
        - POST:
            - Přijímá JSON seznam objektů s detaily žáků.
            - Pro každého žáka:
                - Pokusí se najít odpovídající záznam v databázi (`Zak`) na základě zadaných údajů.
                - Pokud je žák nalezen:
                    - Vytvoří záznam (`Zapsany_zak`) s údaji o termínu, kategorii oprávnění a typu zkoušky.
                - Pokud žák není nalezen:
                    - Přidá do seznamu chyb odpovídající zprávu.
            - Pokud existují chyby (některý žák nebyl nalezen):
                - Vrátí seznam chyb (`400`).
            - Pokud nejsou žádné chyby:
                - Uloží všechny záznamy do databáze (`db.session.commit()`).
                - Vrátí zprávu o úspěchu (`200`).
    """
    try:
        data = request.get_json()
        errors = []

        for student in data:
            evidence_number = student.get('evidence_number')
            first_name = student.get('first_name')
            first_name = ''.join(first_name.split()).capitalize()
            last_name = student.get('last_name')
            last_name = ''.join(last_name.split()).capitalize()
            birth_date = student.get('birth_date')
            license_category = student.get('license_category')
            exam_type = student.get('exam_type').replace('_', ' ') # value se vrací ve tvaru něco_něco tak se mění '_' v ' '
            

            zak = Zak.query.filter_by(ev_cislo=evidence_number, jmeno=first_name, prijmeni=last_name, narozeni=birth_date).first()
            if zak:
                if not zak.konec:
                    errors.append(f'Zak {evidence_number} {first_name} {last_name} nema ukonceny vycvik. Nelze jej zapsat.')
                else:
                    zapis = Zapsany_zak(typ_zkousky=license_category, druh_zkousky=exam_type, id_terminu=session.get('term_id'), id_autoskoly=zak.id_autoskoly,id_zaka=zak.id)
                    db.session.add(zapis)
            else:
                errors.append(f'Zaka {evidence_number} {first_name} {last_name} se nepodarilo nalezt, zkontrolujte si, ze jste informace zadali spravne')
            
        if errors:
            return jsonify({"error": errors}), 400
        
        db.session.commit()
        return jsonify({"message": "Data přijata úspěšně"}), 200 # Odeslání odpovědi o úspěchu
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400

@login_required
@app.route('/enroll', methods=['POST'])
def enroll_drivers():
    """
    Potvrzení účasti studenta na termínu.

    Endpoint nastaví účast studenta na termínu, přiřadí čas začátku zkoušky a komisaře. 
    Vytvoří upozornění pro autoškolu o potvrzení termínu.

    Vrací:
        - `200`: Úspěšné potvrzení.
        - `400`: Chyba při zpracování.

    Metody:
        - POST:
            - JSON požadavek:
                - `id` (int): ID studenta.
                - `time_start` (str): Čas začátku zkoušky.
                - `commissar` (int): ID komisaře.
            - Načte studenta a termín z databáze.
            - Aktualizuje stav a vytvoří záznam o potvrzení.
    """
    try:
        # Získání dat ve formátu JSON
        data = request.get_json()
        
        # Zpracování dat
    
        id_studenta = data.get('id')
        time_start = data.get('time_start')
        id_commissar = data.get('commissar')
        
        zak = Zak.query.filter_by(id= id_studenta).first()
        termin = Termin.query.filter_by(id=session.get('term_id')).first()
        if zak:
            zapis = Zapsany_zak.query.filter_by(id_terminu=session.get('term_id'), id_zaka=id_studenta).first()
            zapis.potvrzeni = 'Y'
            zapis.zacatek = time_start
            zapis.id_komisare = id_commissar
            upozorneni= Upozorneni(zprava= f'Student/ka {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl/a zapsán/a na termín {termin.datum.strftime("%d.%m.%Y")} v {time_start}',
                                   id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            db.session.add(upozorneni)
            db.session.commit()
            return jsonify({"message": "Data přijata úspěšně"}), 200 # Odeslání odpovědi o úspěchu 
        else:
            return jsonify({"error": str(e)}), 400   
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400   

@login_required
@app.route('/api/edit_student', methods=['POST'])
def edit_student():
    if not current_user.isAdmin:
        abort(404)
    try:
        # Načti data z požadavku
        data = request.get_json()
        student_id = data.get('id')
        new_time = data.get('time')
        new_commissar = data.get('commissar')

        zapis = Zapsany_zak.query.filter_by(id_terminu=session.get('term_id'), id_zaka=student_id).first()
        if not zapis:
            print('Nebyl nalezen zápis.')
            return jsonify({"status": "error", "message": "Došlo k chybě při ukládání změn."}), 500
        
        print(new_time != zapis.zacatek, new_time, zapis.zacatek)
        if new_time != zapis.zacatek:
            zapis.zacatek = new_time
            zak = Zak.query.filter_by(id=student_id).first()
            termin = Termin.query.filter_by(id=session.get('term_id')).first()
            upozorneni = Upozorneni(zprava=f"Studentovi/Studentce {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} byl pozměněn začátek na termín {termin.datum.strftime("%d.%m.%Y")} na {new_time}.", id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            db.session.add(upozorneni)

        zapis.id_komisare = new_commissar
        

        db.session.commit()

        print(f"Upravuji studenta {student_id}: nový čas: {new_time}, nový komisař: {new_commissar}")

        # Vrácení úspěšné odpovědi
        return jsonify({"status": "success", "message": "Změny byly uloženy."}), 200
    except Exception as e:
        # Ošetření chyby a vrácení chybové odpovědi
        print(f"Chyba při zpracování změn: {str(e)}")
        return jsonify({"status": "error", "message": "Došlo k chybě při ukládání změn."}), 500

@login_required
@app.route('/api/add_vehicle', defaults={'id': None}, methods=['POST'])
@app.route('/api/add_vehicle/<id>', methods=['POST'])
def add_vehicle(id):
    """
    Přidání vozidla.

    Endpoint přidá nové vozidlo do systému. 

    Vrací:
        - `200`: Přesměrování na profil po úspěšném přidání.
        - `400`: Chyba při zpracování.

    Metody:
        - POST:
            - Přijímá formulářová data:
                - `znacka` (str): Značka vozidla.
                - `model` (str): Model vozidla.
                - `spz` (str): SPZ vozidla.
            - Uloží záznam do databáze.
    """
    id_as = None

    # Kontrola, zda current_user je admin a bylo poskytnuto ID
    if hasattr(current_user, 'isAdmin') and current_user.isAdmin and id:
        id_as = id
    # Pokud není admin a ID není poskytováno, použít current_user.id
    elif not hasattr(current_user, 'isAdmin') or not current_user.isAdmin:
        if id is None:
            id_as = current_user.id
        else:
            return jsonify({"error": "Access denied: Only admins can specify an ID"}), 403
    
    znacka= request.form['znacka']
    model= request.form['model']
    spz= request.form['spz']
    try:
        vozidlo = Vozidlo(znacka=znacka, model=model, spz=spz, id_autoskoly= id_as)
        db.session.add(vozidlo)
        db.session.commit()
        
        return redirect(url_for('profile'))
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400
    
@app.route('/api/delete_student_in_term', methods=['POST'])
@login_required
def delete_student_in_term():
    """
    Vymazaní studenta, který už je zapsán a potvrzen na termínu. 

    Vrací:
        - `200`: Přesměrování na profil po úspěšném přidání.
        - `400`: Chyba při zpracování.

    Metody:
        - POST:
            - Přijímá formulářová data:
                - `id` (int): id žáka.
                - `term_id` (str): id termínu.
            - Vytvoří upozornění pro autoškolu, že žák byl odebrán.
            - Smaže zapsany_zak, který odpovídá id studenta a id termínu.
    """
    if not current_user.isAdmin:
        abort(404)

    data = request.get_json()
    student_id = data.get('id')
    term_id = session.get('term_id')

    if not student_id or not term_id:
        return jsonify({"message": "Missing student ID or tern ID"}), 400
    
    try:
        # Logika pro odstranění studenta z databáze
        zapsany_zak = Zapsany_zak.query.filter_by(id_zaka=student_id, id_terminu=term_id).first()
        zak = Zak.query.filter_by(id=student_id).first()
        termin = Termin.query.filter_by(id= term_id).first()
        if not zapsany_zak:
            return jsonify({"message": "Record not found"}), 404

        formatted_date = termin.datum.strftime("%d.%m.%Y")

        upozorneni = Upozorneni(zprava=f"Student/ka {zak.jmeno} {zak.prijmeni}  byl/a odebrán/a z termínu {formatted_date}.", id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
        db.session.add(upozorneni)
        db.session.delete(zapsany_zak)
        db.session.commit()
        return jsonify({"message": "Student successfully deleted"}), 200
    except Exception as e:
        return jsonify({"message": str(e)}), 500

@login_required
@app.route('/api/delete_student', methods=['POST'])
def delete_student():
    """
    Odebrání žáka z termínu.

    Endpoint umožňuje autoškole odebrat studenta z termínu, pokud jeho účast nebyla potvrzena.

    Vrací:
        - `200`: Úspěšné odebrání studenta.
        - `404`: Student nebo záznam nenalezen.
        - `400`: Chyba při zpracování.

    Metody:
        - POST:
            - Přijímá JSON data:
                - `zak_id` (int): ID studenta.
                - `termin_id` (int): ID termínu.
            - Smaže záznam o účasti z tabulky `Zapsany_zak`.
            - Přidá logovací záznam o odebrání.
    """
    data = request.get_json()
    zak_id = data.get('zak_id')
    termin_id = data.get('termin_id')
    
    # Najdi záznam v tabulce Zapsany_zak a smaž ho
    zapsany = Zapsany_zak.query.filter_by(id_zaka=zak_id, id_terminu=termin_id).filter(Zapsany_zak.potvrzeni.in_(['W', 'N'])).first()
    zak = Zak.query.filter_by(id=zak_id).first()
    zaznam = Zaznam(druh='odpis', kdy=datetime.now(),
                    zprava=f'Autoškola odebrala studentku/studenta {zak.jmeno} {zak.prijmeni} {zak.ev_cislo} z termínu s id: {session.get('term_id')}',
                    id_autoskoly=current_user.id)
    db.session.add(zaznam)

    if zapsany:
        db.session.delete(zapsany)
        db.session.commit()
        return jsonify({'message': 'Student smazán'}), 200
    else:
        print(f'Nebylo nalezeno zapsání {zak_id} {termin_id}')
        return jsonify({'error': 'Student nenalezen'}), 404

@login_required
@app.route('/api/sign_up', methods=['POST'])
def docx_for_signup():
    """
    API metoda, která vytvoří dokument o zápisu studentů. Zapsané studenty přidá do databáze. 

    Parametry:
        request: JSON soubor s novými studenty, adresou učebny, typem výuky, vozidli pro výuku

    Vrací:
        str: pokud při vytvoření a nebo commitu do databáze vznikne error
        str: pokud vše proběhne v pořádku
        docx: dokument o žádost zápisu studentů do výuky a výcviku
    """
    try:
        data = request.get_json() #data z POST requestu

        driving_school_id = data['main_form'].get('driving_school_id')

        if driving_school_id:
            autoskola = Autoskola.query.filter_by(id=data['main_form']['driving_school_id']).first()
        else:
            autoskola = Autoskola.query.filter_by(id=current_user.id).first() # data o autoškole, potřebuju jméno pro ukládání

        adresa = data['main_form']['adress'] # adresa účebny
        datum = data['main_form']['start_of_training']
        datum = datetime.strptime(datum, "%Y-%m-%d") # datum začátku výcviku
        seznam_vozidel = data['main_form']['vehicle_list'] # seznam vozidel k výcviku
        seznam_studentu = data['students'] # seznam studentů do výcviku

        document = Document() # docx dokument
        document.sections[0].right_margin = 457200
        document.sections[0].left_margin = 457200
        document.add_heading('Seznam řidičů žádajících o zařazení do výuky a výcviku', 0) # nadpis

        document.add_paragraph(f'Adresa učebny: {adresa}')
        document.add_paragraph(f'Datum začátku výcviku: {datum.strftime("%d.%m.%Y")}')

        document.add_heading('Seznam vozidel k výcviku:', level=1)
        for vozidlo in seznam_vozidel:
            document.add_paragraph(vozidlo)

        save_document = False
        duplicate_students = []

        # vytvoření tabulky a sloupců s názvem
        table = document.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Evidenčí číslo'
        hdr_cells[1].text = 'Jméno'
        hdr_cells[2].text = 'Přijmení'
        hdr_cells[3].text = 'Datum narození'
        hdr_cells[4].text = 'Adresa'
        hdr_cells[5].text = 'Číslo řidičského průkazu'
        hdr_cells[6].text = 'Skupina řidičského oprávnění'
        hdr_cells[7].text = 'Druh výcviku'

        for student in seznam_studentu: # cykl pro studenty, aby se zapsali do tabulky a zároveň jejich tvorba do db
            exist_student = Zak.query.filter_by(ev_cislo=student['evidence_number'], id_autoskoly= autoskola.id).first()
            typ_vycviku = student['type_of_teaching'].replace('-', ' ')
            podminka = typ_vycviku in ['opakovaná výuka', 'opakovaný výcvik'] # pokud je žák zapsaný na opakovaný výcvik nebo výuku tak se nemá znovu vytvářet ale ani nevracet jako chybný vstup
            ridicky_prukaz = student['drivers_license'] if student['drivers_license'] else None
            prvni_pokus = student['first_try'] if student['first_try'] else None
            if exist_student and not podminka:
                duplicate_students.append(student['evidence_number'])
                continue  # Nepřidáváme duplicity
            
            save_document = True
            row_cells = table.add_row().cells
            row_cells[0].text = student['evidence_number']
            row_cells[1].text = student['first_name']
            row_cells[2].text = student['last_name']
            row_cells[3].text = student['birth_date']
            row_cells[4].text = student['adress']
            row_cells[5].text = ridicky_prukaz if ridicky_prukaz else ' '
            row_cells[6].text = typ_vycviku
            row_cells[7].text = student['license_category']

            # Tady se vytvoří žák a uloží se do db
            if (exist_student and podminka):
                zaznam = Zaznam(druh='přidání', kdy=datetime.now(), zprava=f'Autoškola zapsala studentku/studenta {student['first_name']} {student['last_name']} {student['evidence_number']} na opravnou výuku nebo výcvik.', id_autoskoly=autoskola.id)
            else:
                zak = Zak(ev_cislo=student['evidence_number'], jmeno=student['first_name'], prijmeni=student['last_name'], narozeni=student['birth_date'], adresa=student['adress'], id_autoskoly= autoskola.id, zacatek=datum, cislo=ridicky_prukaz, prvni=prvni_pokus)
                zaznam = Zaznam(druh='přidání', kdy=datetime.now(), zprava=f'Autoškola zapsala studentku/studenta {student['first_name']} {student['last_name']} {student['evidence_number']} do výuky a výcviku.', id_autoskoly=autoskola.id)
                db.session.add(zaznam)
                db.session.add(zak)
                
        db.session.commit()
        document.add_page_break()

        # Název složky a souboru
        folder_path = "Zapis_studentu"
        file_name = f"{autoskola.nazev.replace(" ", "_")}_{datum.strftime("%d-%m-%Y")}.docx"

        # Kontrola a vytvoření složky
        Path(folder_path).mkdir(parents=True, exist_ok=True)

        if save_document:
            document.save(f"{folder_path}/{file_name}")
        
        if duplicate_students:
            return jsonify({
                "error": "Někteří studenti již existují v systému.",
                "duplicate_evidence_numbers": duplicate_students
            }), 400

        return jsonify({"message": "Data přijata úspěšně"}), 200
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400

@login_required
@app.route('/api/make_sheet', methods=['GET'])
def make_sheet():
    """
    Vytvoří a vrátí Word dokument se seznamem zapsaných žáků na daný termín.

    Funkce načte ID termínu ze session, dohledá odpovídající termín v databázi
    a získá seznam přihlášených žáků. Následně vygeneruje Word dokument, uloží
    ho do paměti a odešle jako přílohu ke stažení.

    Návratová hodnota:
        - Pokud je ID termínu neplatné nebo termín neexistuje, vrátí HTTP 400.
        - Pokud nejsou k dispozici žádní zapsaní žáci, může vrátit prázdný dokument.
        - Jinak vrátí vygenerovaný Word dokument jako přílohu ke stažení.
    """

    # Načtení ID termínu ze session
    term_id = session.get('term_id')
    if not term_id:
        abort(400, "ID termínu není k dispozici.")

    # Načtení termínu z databáze
    term = Termin.query.filter_by(id=term_id).first()
    if not term:
        abort(400, "Termín s tímto ID neexistuje.")

    # Načtení zapsaných žáků pro daný termín
    zaci = Zapsany_zak.query \
                      .join(Zak) \
                      .join(Termin) \
                      .filter(Zapsany_zak.id_terminu == term_id) \
                      .all()

    if not zaci:
        abort(400, "Pro tento termín nejsou zapsaní žádní žáci.")

    # Vytvoření dokumentu
    dokument = app_logic.create_student_document(term.datum, zaci)
    if not dokument:
        abort(500, "Nepodařilo se vygenerovat dokument.")

    # Uložení dokumentu do paměti a odeslání ke stažení
    file_stream = BytesIO()
    dokument.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, 
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                     as_attachment=True, 
                     download_name=f'{term.datum}_seznam_studentu.docx')

@login_required
@app.route('/api/generate_doc', methods=['POST'])
def generate_doc():
    """
    Generování dokumentu.

    Endpoint vytvoří a vygeneruje dokument obsahující informace o termínu a zapsaných žácích.

    Vrací:
        - Soubor ke stažení (`.docx`) s vygenerovaným dokumentem.
        - `400`: Chyba při zpracování požadavku nebo chybějící data.

    Metody:
        - POST:
            - Přijímá JSON data:
                - `name` (str): Název autoškoly.
                - `date` (str): Datum termínu (formát `YYYY-MM-DD`).
            - Načte autoškolu, termín, zapsané žáky a komisaře.
            - Vygeneruje dokument pomocí `app_logic.create_document`.
    """
    data = request.get_json()
    name = data.get('name')
    date = data.get('date')

    # Zkontrolujeme, zda jsme obdrželi data
    if not name or not date:
        return jsonify({"error": "Název a datum jsou povinné"}), 400
    autoskola = Autoskola.query.filter_by(nazev = name).first()
    datum = datetime.strptime(date, "%Y-%m-%d")

    zaci = Zapsany_zak.query \
                    .join(Zak) \
                    .join(Termin) \
                    .filter(Termin.datum == date, Zak.id_autoskoly == autoskola.id, Zapsany_zak.potvrzeni == 'Y') \
                    .all()
    
    if int(current_user.id) > 1000000:
        komisar = Superadmin.query.filter_by(id=current_user.id).first()
    else:
        komisar = Komisar.query.filter_by(id=current_user.id).first()

    dokument = app_logic.create_document(autoskola, datum, zaci, komisar)

    # Uložíme dokument do paměti
    file_stream = BytesIO()
    dokument.save(file_stream)
    file_stream.seek(0)

    # Odpovíme stažením souboru
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{name}_{date}_document.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@login_required
@app.route('/api/success', methods=['POST'])
def student_success():
    """
    API metoda, která slouží pro zápis úspěchu studenta na zkoušce. 

    Parametry:
        request: JSON soubor s id studenta, který uspěl.

    Vrací:
        str: error, když dojede k chybě nebo když neexistuje termín nebo žák
        str: pokud vše proběhne v pořádku
    """
    if not current_user.isCommissar:
        abort(404)
    try:
        # Získání dat ve formátu JSON
        data = request.get_json()
        
        # Zpracování dat
        id_studenta = data.get('id')
        
        zak = Zak.query.filter_by(id = id_studenta).first()
        termin = Termin.query.filter_by(id = session.get('term_id')).first()
        if zak and termin:
            zapis = Zapsany_zak.query.filter_by(id_terminu=session.get('term_id'), id_zaka=id_studenta).first()
            zapis.zaver = 'Y'
            zak.splnil = termin.datum
            upozorneni= Upozorneni(zprava= f'Studentka/Student {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} úspěšně splnil termín.', id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            db.session.add(zapis)
            db.session.add(upozorneni)
            db.session.commit()
            
            return jsonify({"message": "Data přijata úspěšně"}), 200 # Odeslání odpovědi o úspěchu 
        else:
            return jsonify({"error": str(e)}), 400   
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400    

@login_required
@app.route('/api/reject', methods=['POST'])
def student_reject():
    """
    API metoda, která slouží k zápisu neúspěchu při zkoušce. 

    Parametry:
        request: JSON soubor s id žáka

    Vrací:
        str: error, pokud se nenajde žák nebo termín nebo pokud dojde k nějaké chybě
        str: 200, pokud vše proběhne v pořádku
        abort: 404 pokud user není admin, jelikož k metodě by měl mít přístup pouze admin
    """
    if not current_user.isCommissar:
        abort(404)
    try:
        # Získání dat ve formátu JSON
        data = request.get_json()
        
        # Zpracování dat
        id_studenta = data.get('id')
        
        zak = Zak.query.filter_by(id= id_studenta).first()
        termin = Termin.query.filter_by(id = session.get('term_id')).first()
        if zak and termin:
            zapis = Zapsany_zak.query.filter_by(id_terminu=session.get('term_id'), id_zaka=id_studenta).first()
            zapis.zaver = 'N'
            if zak.prvni == None:
                zak.prvni = termin.datum
            upozorneni= Upozorneni(zprava= f'Studentka/Student {zak.ev_cislo} {zak.jmeno} {zak.prijmeni} neuspěl u termínu.', id_autoskoly= zak.id_autoskoly, datum_vytvoreni=datetime.now())
            db.session.add(upozorneni)
            db.session.add(zapis)
            db.session.commit()
            
            return jsonify({"message": "Data přijata úspěšně"}), 200 # Odeslání odpovědi o úspěchu 
        else:
            return jsonify({"error": str(e)}), 400   
    except Exception as e:
        # Pokud nastane chyba, odeslat chybovou zprávu
        return jsonify({"error": str(e)}), 400    

@login_required
@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    """
    Získání upozornění.

    Endpoint vrací seznam posledních 10 upozornění pro aktuální autoškolu, 
    seřazených podle data vytvoření (od nejnovějšího).

    Vrací:
        - `200`: Seznam upozornění (bez nových zpráv).
        - `201`: Seznam upozornění (s novými zprávami).
        - `400`: Chyba při zpracování požadavku.

    Metody:
        - GET:
            - Načte záznamy z tabulky `Upozorneni`:
                - `zprava` (str): Text zprávy.
                - `stav` (str): Stav zprávy (`N` = nová, `R` = přečtená).
            - Vrací seznam upozornění jako JSON.
    """
    try:
        upozorneni_list = Upozorneni.query.filter_by(id_autoskoly=current_user.id)\
                                      .order_by(Upozorneni.datum_vytvoreni.desc())\
                                      .limit(10)\
                                      .all()
        
        lst = []
        new_notif_fount = False # byla nalezeno nové upozornění?

        for upozorneni in upozorneni_list:
            lst.append({
            'zprava': upozorneni.zprava,
            'stav': upozorneni.stav
            })
            if upozorneni.stav == 'N':
                new_notif_fount = True # pokud je nová zpráva zapíše se   
        
        if new_notif_fount:
            return jsonify(lst), 201
        
        return jsonify(lst), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 400

@login_required
@app.route('/api/notifications_change', methods=['POST'])
def change_notifications():
    """
    Změna stavu upozornění.

    Endpoint nastaví stav posledních 10 upozornění aktuální autoškoly na "přečtené" (`stav = 'Y'`).

    Vrací:
        - `204`: Úspěšná změna stavu.
        - `400`: Chyba při zpracování požadavku.

    Metody:
        - POST:
            - Načte posledních 10 upozornění z tabulky `Upozorneni`.
            - Změní jejich stav na `Y` (přečteno).
            - Uloží změny do databáze.
    """
    try:
        upozorneni_list = Upozorneni.query.filter_by(id_autoskoly=current_user.id)\
                                      .order_by(Upozorneni.datum_vytvoreni.desc())\
                                      .limit(10)\
                                      .all()

        for upozorneni in upozorneni_list:
            upozorneni.stav = 'Y'
        
        db.session.commit()    
        
        
        return '', 204
    
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@login_required
@app.route('/api/get_vehicles', methods=['GET'])
def get_vehicles():
    """
    API metoda, která vrací seznam vozidel jednotlivých autoškol. Podle poskytnutého id si najde vozidla s ním spojená a pošle je zpět. 

    Parametry:
        request: JSON soubor s novými studenty, adresou učebny, typem výuky, vozidli pro výuku

    Vrací:
        str: pokud při vytvoření a nebo commitu do databáze vznikne error
        str: pokud vše proběhne v pořádku
        docx: dokument o žádost zápisu studentů do výuky a výcviku
    """
    autoskola = request.args.get('value')
    vozidla = Vozidlo.query.filter_by(id_autoskoly=autoskola).all()
    ls_vozidla = []
    for vozidlo in vozidla:
        ls_vozidla.append({'id': vozidlo.id,
                           'znacka':vozidlo.znacka,
                           'model': vozidlo.model,
                           'RZ': vozidlo.spz
                          })
    return jsonify(ls_vozidla)

@login_required
@app.route('/api/change_info_ds', methods=['POST'])
def change_info():
    """
        API endpoint pro změnu informací o AŠ na profilu.
        #TODO
    """
    try:
        # Načti JSON data z požadavku
        data = request.get_json()

        if not data:
            return jsonify({"error": "No data provided"}), 400

        # Najdi první záznam autoškoly (předpoklad, že je jedna autoškola, případně modifikujte)
        autoskola = Autoskola.query.first()
        if not autoskola:
            return jsonify({"error": "Autoskola not found"}), 404

        # Aktualizuj pouze poskytnutá pole
        if 'name' in data:
            autoskola.nazev = data['name']
        if 'email' in data:
            autoskola.email = data['email']
        if 'adress' in data:
            autoskola.adresa_u = data['adress']
        if 'datovka' in data:
            autoskola.da_schranka = data['datovka']

        # Ulož změny do databáze
        db.session.commit()

        return jsonify({"message": "Autoskola updated successfully"}), 200
    except Exception as e:
        return jsonify({"error": "An unexpected error occurred", "details": str(e)}), 500

@app.errorhandler(404)
def error404(e):
    return render_template('404.html'), 404

@app.route('/logout')
def logout():
    #TODO dokumentace
    if current_user.is_authenticated:
        if not current_user.isCommissar:
            zaznam = Zaznam(druh='odhlásil se', kdy=datetime.now(), zprava='Autoškola se odhlásila z aplikace', id_autoskoly=current_user.id)
            db.session.add(zaznam)
            db.session.commit()
        session.clear()
        logout_user()
    return redirect(url_for('home'))

@app.context_processor
def inject_globals():
    return {
        'version': '0.9.9'
    }

@app.before_request # tahle metoda se spustí před každým requestem, brání v prodloužení sessionu pro def get_notification():
def dont_extend_session():
    # Zkontrolujte, zda je požadavek na specifický endpoint
    if request.endpoint == '/api/notifications' and current_user.is_authenticated:
        session.permanent = False  # Dočasně vypneme permanentní session pro tento endpoint

@loginManager.user_loader
def load_user(user_id):
    """
    Callback funkce pro Flask-Login, která načte uživatele na základě jeho ID.

    Tato funkce se používá Flask-Login rozšířením při každém požadavku k identifikaci 
    aktuálně přihlášeného uživatele. Načítá uživatele z aplikační logiky pomocí jeho ID 
    a volitelně nastaví, zda je uživatel administrátorem na základě session proměnné.

    Args:
        user_id (str): ID uživatele, které je uloženo v session.

    Returns:
        User: Instance třídy User s informací o ID a oprávnění správce.
    """
    is_admin = session.get('isAdmin', False)
    return app_logic.User(user_id, is_admin)

if __name__ == "__main__":
    app.run(debug=True)